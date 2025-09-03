"""
Main application file for the Medical Laboratory Management System
"""
import tkinter as tk
from tkinter import ttk, messagebox
import tkinter.filedialog as filedialog
import tkinter.font as tkFont
import sqlite3
from datetime import datetime
from typing import List, Optional
import hashlib
import uuid
from docx import Document
from database import DatabaseManager
from models import (
    Patient, TestType, TestRequest, Sample, MedicalReport, 
    Invoice, User, InventoryItem, PurchaseOrder, TestTemplate, Gender, 
    TestStatus, SampleStatus, UserRole, PaymentMethod, Permission, UserPermission
)
from utils import generate_barcode, send_email, encrypt_data, decrypt_data
from translations import _, set_language, register_language_change_callback

from translations import _, set_language, register_language_change_callback

class MedicalLabApp:
    def __init__(self, root):
        self.root = root
        self.root.title(_("Medical Laboratory Management System"))
        self.root.geometry("1200x800")
        
        # Configure 3D style
        self.configure_3d_style()
        
        # Initialize database
        self.db = DatabaseManager()
        
        # Current user
        self.current_user = None
        
        # Register for language change notifications
        register_language_change_callback(self.on_language_change)
        
        # Setup UI
        self.setup_ui()
        
        # Load initial data
        self.load_initial_data()
    
    def configure_3d_style(self):
        """Configure 3D style with beautiful colors for the application"""
        style = ttk.Style()
        
        # Configure improved color scheme for better visibility with black text
        bg_color = "#f5f7fa"  # Light gray-blue background
        accent_color = "#3498db"  # Bright blue accent
        accent_hover = "#2980b9"  # Darker blue for hover
        button_color = "#3498db"  # Blue for buttons
        button_hover = "#2980b9"  # Darker blue for button hover
        text_color = "#000000"  # Black text color for maximum clarity
        header_color = "#2980b9"  # Darker blue header color
        nav_color = "#3498db"  # Blue navigation color
        card_bg = "#ffffff"  # White card background
        
        # Configure styles
        style.configure("TFrame", background=bg_color)
        style.configure("TLabel", background=bg_color, foreground=text_color)
        style.configure("Header.TFrame", background=header_color)
        style.configure("Nav.TFrame", background=nav_color)
        style.configure("Card.TFrame", background=card_bg, relief="raised", borderwidth=4)
        style.configure("Title.TLabel", background=bg_color, foreground=text_color, 
                       font=("Arial", 18, "bold"))
        style.configure("Header.TLabel", background=header_color, foreground="#000080", 
                       font=("Arial", 14, "bold"))
        style.configure("Nav.TLabel", background=nav_color, foreground="#000080")
        
        # Button styles with enhanced 3D effect
        style.configure("TButton", 
                       background=button_color, 
                       foreground="#000080",
                       borderwidth=5,
                       relief="raised",
                       font=("Arial", 11, "bold"))
        style.map("TButton",
                 background=[("active", button_hover)],
                 relief=[("pressed", "sunken")])
        
        # Navigation button style
        style.configure("Nav.TButton", 
                       background=nav_color, 
                       foreground="black",
                       borderwidth=4,
                       relief="raised",
                       font=("Arial", 10, "bold"))
        style.map("Nav.TButton",
                 background=[("active", "#2c3e50")],
                 relief=[("pressed", "sunken")])
        
        # Accent button style
        style.configure("Accent.TButton", 
                       background="#ff6600",  # Orange color for better visibility
                       foreground="#000080",
                       borderwidth=5,
                       relief="raised",
                       font=("Arial", 11, "bold"))
        style.map("Accent.TButton",
                 background=[("active", "#cc5200")],  # Darker orange on hover
                 relief=[("pressed", "sunken")])
        
        # Configure root window background
        self.root.configure(bg=bg_color)
    
    def setup_ui(self):
        # Create main frames with 3D styling
        self.header_frame = ttk.Frame(self.root, style="Header.TFrame")
        self.header_frame.pack(fill=tk.X, padx=5, pady=5)
        
        self.nav_frame = ttk.Frame(self.root, style="Nav.TFrame")
        self.nav_frame.pack(fill=tk.Y, side=tk.LEFT, padx=5, pady=5)
        
        self.content_frame = ttk.Frame(self.root, style="Card.TFrame")
        self.content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.footer_frame = ttk.Frame(self.root, style="Card.TFrame")
        self.footer_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Header with login info and language selector
        self.setup_header()
        
        # Navigation
        self.setup_navigation()
        
        # Initial content
        self.show_login_screen()
    
    def setup_header(self):
        # Language selector
        lang_frame = ttk.Frame(self.header_frame, style="Header.TFrame")
        lang_frame.pack(side=tk.RIGHT, padx=10, pady=5)
        
        ttk.Label(lang_frame, text=_("Language:"), style="Header.TLabel").pack(side=tk.LEFT)
        
        self.lang_var = tk.StringVar(value="en")
        lang_combo = ttk.Combobox(lang_frame, textvariable=self.lang_var, 
                                 values=["en", "ar"], state="readonly", width=5)
        lang_combo.pack(side=tk.LEFT, padx=5)
        lang_combo.bind("<<ComboboxSelected>>", self.change_language)
        
        # User info and controls
        user_controls_frame = ttk.Frame(self.header_frame, style="Header.TFrame")
        user_controls_frame.pack(side=tk.RIGHT, padx=20, pady=5)
        
        # System name edit icon (only visible when logged in as admin)
        self.edit_system_name_btn = tk.Button(user_controls_frame, text=_("✏️ System Name"), 
                                             command=self.edit_system_name, state=tk.DISABLED,
                                             bg="#3498db", fg="#000080",  # Blue background with dark blue text
                                             font=("Arial", 11, "bold"),
                                             relief="raised", bd=4)
        self.edit_system_name_btn.pack(side=tk.RIGHT, padx=10)
        
        # Server connection icon (only visible when logged in as admin)
        self.server_icon_btn = tk.Button(user_controls_frame, text=_("Server"), 
                                        command=self.show_server_connection, state=tk.DISABLED,
                                        bg="#3498db", fg="#000080",  # Blue background with dark blue text
                                        font=("Arial", 11, "bold"),
                                        relief="raised", bd=4)
        self.server_icon_btn.pack(side=tk.RIGHT, padx=10)
        
        # Change password button (only visible when logged in as admin) with red text
        self.change_password_btn = tk.Button(user_controls_frame, text=_("Change Password"), 
                                             command=self.change_admin_password, state=tk.DISABLED,
                                             bg="#3498db", fg="red",  # Blue background with red text
                                             font=("Arial", 11, "bold"),
                                             relief="raised", bd=4)
        self.change_password_btn.pack(side=tk.RIGHT, padx=10)
        
        # User info
        self.user_label = ttk.Label(user_controls_frame, text=_("Not logged in"), style="Header.TLabel")
        self.user_label.pack(side=tk.RIGHT, padx=10)
        
        # Logout button with blue text
        self.logout_btn = tk.Button(user_controls_frame, text=_("Logout"), 
                                    command=self.logout, state=tk.DISABLED,
                                    bg="#3498db", fg="blue",  # Blue background with blue text
                                    font=("Arial", 11, "bold"),
                                    relief="raised", bd=4)
        self.logout_btn.pack(side=tk.RIGHT)
        
        # Application title
        title_frame = ttk.Frame(self.header_frame, style="Header.TFrame")
        title_frame.pack(side=tk.LEFT, padx=20, pady=5)
        
        self.title_label = ttk.Label(title_frame, text=_("Medical Laboratory Management System"), 
                                    style="Header.TLabel", font=("Arial", 14, "bold"))
        self.title_label.pack()
    
    def setup_navigation(self):
        # Navigation header
        self.nav_header_label = ttk.Label(self.nav_frame, text=_("Navigation"), 
                                         style="Nav.TLabel", font=("Arial", 12, "bold"))
        self.nav_header_label.pack(fill=tk.X, padx=5, pady=5)
        
        # Navigation buttons (will be enabled after login)
        self.nav_buttons = {}
        nav_items = [
            (_("Dashboard"), self.show_dashboard),
            (_("Patients"), self.show_patients),
            (_("Tests"), self.show_tests),
            (_("Samples"), self.show_samples),
            (_("Reports"), self.show_reports),
            (_("Billing"), self.show_billing),
            (_("Inventory"), self.show_inventory),
            (_("Users"), self.show_users),
            (_("Statistics"), self.show_statistics)
        ]
        
        for i, (text, command) in enumerate(nav_items):
            btn = ttk.Button(self.nav_frame, text=text, command=command, 
                            width=15, state=tk.DISABLED, style="Nav.TButton")
            btn.pack(fill=tk.X, pady=3, padx=5)
            self.nav_buttons[text] = btn
    
    def on_language_change(self):
        """Callback function called when language changes"""
        # Update all UI elements that need translation
        self.update_ui_texts()
        
        # Refresh current screen to update any dynamically created elements
        if hasattr(self, 'current_screen') and self.current_screen:
            self.current_screen()
    
    def update_ui_texts(self):
        """Update all UI texts when language changes"""
        # Update window title
        self.root.title(_("Medical Laboratory Management System"))
        
        # Update header elements
        self.title_label.config(text=_("Medical Laboratory Management System"))
        
        # Update language selector label
        # Note: We can't directly update the label text in the combobox, but we can update other labels
        
        # Update user info
        if self.current_user:
            self.user_label.config(text=_("Logged in as: {} ({})").format(
                self.current_user.username, _(self.current_user.role.value)))
        else:
            self.user_label.config(text=_("Not logged in"))
        
        # Update button texts
        self.logout_btn.config(text=_("Logout"))
        self.change_password_btn.config(text=_("Change Password"))
        self.server_icon_btn.config(text=_("Server"))
        
        # Update navigation header
        self.nav_header_label.config(text=_("Navigation"))
        
        # Update navigation buttons
        nav_items = [
            _("Dashboard"),
            _("Patients"),
            _("Tests"),
            _("Samples"),
            _("Reports"),
            _("Billing"),
            _("Inventory"),
            _("Users"),
            _("Statistics")
        ]
        
        for i, text in enumerate(nav_items):
            btn = list(self.nav_buttons.values())[i]
            btn.config(text=text)
    
    def change_language(self, event=None):
        """Handle language change"""
        # Change the language
        lang = self.lang_var.get()
        set_language(lang)
        
        # The on_language_change callback will be automatically called
        # which will update all UI elements
    
    def load_initial_data(self):
        # Create admin user if not exists
        admin_user = self.db.get_user_by_username("admin")
        if not admin_user:
            admin = User(
                id=str(uuid.uuid4()),
                username="admin",
                email="admin@lab.com",
                password_hash=self.hash_password("admin123"),
                role=UserRole.ADMIN
            )
            self.db.create_user(admin)
        
        # Create some sample test types if none exist
        test_types = self.db.get_all_test_types()
        if not test_types:
            sample_tests = [
                TestType(str(uuid.uuid4()), _("Complete Blood Count"), 
                        _("Measures several components and features of blood"), 50.0, _("Blood")),
                TestType(str(uuid.uuid4()), _("Urinalysis"), 
                        _("Analysis of urine components"), 30.0, _("Urine")),
                TestType(str(uuid.uuid4()), _("Stool Analysis"), 
                        _("Examination of stool sample"), 40.0, _("Stool")),
                TestType(str(uuid.uuid4()), _("X-Ray Chest"), 
                        _("Chest X-Ray imaging"), 100.0, _("Radiology"))
            ]
            
            for test in sample_tests:
                self.db.create_test_type(test)
    
    def hash_password(self, password: str) -> str:
        return hashlib.sha256(password.encode()).hexdigest()
    
    def show_login_screen(self):
        self.current_screen = self.show_login_screen
        # Clear content frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()
        
        # Login form with 3D styling
        login_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        login_frame.pack(expand=True, padx=20, pady=20)
        
        # Add a title with better styling
        title_label = ttk.Label(login_frame, text=_("Login"), 
                               style="Title.TLabel")
        title_label.pack(pady=20)
        
        # Username field
        username_frame = ttk.Frame(login_frame, style="Card.TFrame")
        username_frame.pack(fill=tk.X, padx=40, pady=10)
        
        ttk.Label(username_frame, text=_("Username:"), font=("Arial", 11), foreground="#2c3e50").pack(anchor=tk.W)
        self.username_entry = ttk.Entry(username_frame, font=("Arial", 11))
        self.username_entry.pack(fill=tk.X, pady=5)
        
        # Password field
        password_frame = ttk.Frame(login_frame, style="Card.TFrame")
        password_frame.pack(fill=tk.X, padx=40, pady=10)
        
        ttk.Label(password_frame, text=_("Password:"), font=("Arial", 11), foreground="#2c3e50").pack(anchor=tk.W)
        self.password_entry = ttk.Entry(password_frame, show="*", font=("Arial", 11))
        self.password_entry.pack(fill=tk.X, pady=5)
        
        # Login button with accent styling
        login_btn = ttk.Button(login_frame, text=_("Login"), 
                              command=self.login, style="Accent.TButton")
        login_btn.pack(pady=20)
        
        # Removed demo login info per user request
        
        # Focus on username entry
        self.username_entry.focus()
    
    def login(self):
        username = self.username_entry.get()
        password = self.password_entry.get()
        
        if not username or not password:
            messagebox.showerror(_("Error"), _("Please enter both username and password"))
            return
        
        hashed_password = self.hash_password(password)
        user = self.db.authenticate_user(username, hashed_password)
        
        if user:
            self.current_user = user
            self.user_label.config(text=_("Logged in as: {} ({})").format(
                user.username, _(user.role.value)))
            self.logout_btn["state"] = tk.NORMAL
            
            # Enable admin-only buttons for admin users
            if user.role == UserRole.ADMIN:
                self.change_password_btn["state"] = tk.NORMAL
                self.server_icon_btn["state"] = tk.NORMAL
                self.edit_system_name_btn["state"] = tk.NORMAL
            else:
                self.change_password_btn["state"] = tk.DISABLED
                self.server_icon_btn["state"] = tk.DISABLED
                self.edit_system_name_btn["state"] = tk.DISABLED
            
            # Enable navigation buttons based on role
            self.enable_navigation()
            
            # Show dashboard
            self.show_dashboard()
        else:
            messagebox.showerror(_("Error"), _("Invalid username or password"))
    
    def logout(self):
        self.current_user = None
        self.user_label.config(text=_("Not logged in"))
        self.logout_btn["state"] = tk.DISABLED
        self.change_password_btn["state"] = tk.DISABLED
        self.server_icon_btn["state"] = tk.DISABLED
        self.edit_system_name_btn["state"] = tk.DISABLED
        
        # Disable navigation
        for btn in self.nav_buttons.values():
            btn.config(state=tk.DISABLED)
        
        # Show login screen
        self.show_login_screen()
    
    def enable_navigation(self):
        # Enable all buttons for admin
        if self.current_user.role == UserRole.ADMIN:
            for btn in self.nav_buttons.values():
                btn.config(state=tk.NORMAL)
            # Enable admin-only buttons for admin users
            self.change_password_btn["state"] = tk.NORMAL
            self.server_icon_btn["state"] = tk.NORMAL
            self.edit_system_name_btn["state"] = tk.NORMAL
        else:
            # Enable based on role (simplified for demo)
            for btn in self.nav_buttons.values():
                btn.config(state=tk.NORMAL)
            # Disable admin-only buttons for non-admin users
            self.change_password_btn["state"] = tk.DISABLED
            self.server_icon_btn["state"] = tk.DISABLED
            self.edit_system_name_btn["state"] = tk.DISABLED
    def show_dashboard(self):
        self.current_screen = self.show_dashboard
        self.clear_content()
        
        # Professional dashboard title
        title_label = ttk.Label(self.content_frame, text=_("🏥 Medical Laboratory Dashboard"), 
                               style="Title.TLabel", font=("Arial", 18, "bold"))
        title_label.pack(pady=15)
        
        # Stats cards with enhanced 3D effect and professional styling
        stats_frame = ttk.LabelFrame(self.content_frame, text=_("📊 Laboratory Statistics"), 
                                    style="Card.TFrame", padding=15)
        stats_frame.pack(fill=tk.X, padx=15, pady=15)
        
        # Real stats from database (in a real app, these would come from the database)
        total_patients = len(self.db.get_all_patients())
        pending_tests = len([tr for tr in self.db.get_all_test_requests() if tr.status == TestStatus.PENDING])
        completed_today = len([mr for mr in self.db.get_all_medical_reports() 
                              if mr.created_at.date() == datetime.now().date()])
        low_inventory = len(self.db.get_low_stock_items())
        
        stats = [
            (_("👥 Total Patients"), str(total_patients)),
            (_("⏳ Pending Tests"), str(pending_tests)),
            (_("✅ Completed Today"), str(completed_today)),
            (_("⚠️ Low Inventory"), str(low_inventory))
        ]
        
        for i, (label, value) in enumerate(stats):
            # Create card with enhanced 3D effect
            card = ttk.Frame(stats_frame, style="Card.TFrame")
            card.grid(row=0, column=i, padx=15, pady=15, sticky="ew")
            stats_frame.grid_columnconfigure(i, weight=1)
            
            # Add 3D effect with shadow
            card_shadow = tk.Frame(stats_frame, bg="#d0d0d0", width=2, height=2)
            card_shadow.grid(row=0, column=i, padx=(17, 13), pady=(17, 13), sticky="se")
            
            ttk.Label(card, text=label, font=("Arial", 11, "bold"), foreground="#000000").pack(pady=12)
            ttk.Label(card, text=value, font=("Arial", 18, "bold"), 
                     foreground="#3498db").pack(pady=7)
        
        # Professional Results management section
        results_frame = ttk.LabelFrame(self.content_frame, text=_("📋 Medical Results Management"),
                                      style="Card.TFrame", padding=15)
        results_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        # Buttons for result actions with professional styling and icons
        button_frame = ttk.Frame(results_frame, style="Card.TFrame")
        button_frame.pack(fill=tk.X, padx=5, pady=15)
        
        ttk.Button(button_frame, text=_("📋 View All Results"), 
                  command=self.show_results, style="Accent.TButton").pack(side=tk.LEFT, padx=8)
        ttk.Button(button_frame, text=_("➕ Create New Result"), 
                  command=self.create_new_result, style="Accent.TButton").pack(side=tk.LEFT, padx=8)
        ttk.Button(button_frame, text=_("📄 Manage Templates"), 
                  command=self.manage_test_templates, style="Accent.TButton").pack(side=tk.LEFT, padx=8)
        ttk.Button(button_frame, text=_("🖨️ Print Results"), 
                  command=self.print_results, style="Accent.TButton").pack(side=tk.LEFT, padx=8)
        
        # Recent results table with enhanced styling
        recent_frame = ttk.LabelFrame(results_frame, text=_("🕒 Recent Results"), 
                                     style="Card.TFrame", padding=10)
        recent_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=10)
        
        # Create treeview for recent results with professional styling
        columns = (_("Result ID"), _("Patient"), _("Test Type"), _("Status"), _("Created At"))
        self.results_tree = ttk.Treeview(recent_frame, columns=columns, show="headings", 
                                        style="Treeview", height=8)
        
        # Configure column headings with better styling
        for col in columns:
            self.results_tree.heading(col, text=col)
            self.results_tree.column(col, width=130)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(recent_frame, orient=tk.VERTICAL, 
                                 command=self.results_tree.yview)
        self.results_tree.configure(yscroll=scrollbar.set)
        
        self.results_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load results data
        self.load_results_data()

    def load_results_data(self):
        # Check if results_tree exists
        if not hasattr(self, 'results_tree'):
            return
            
        # Clear existing data
        for item in self.results_tree.get_children():
            self.results_tree.delete(item)
        
        # Load medical reports from database
        reports = self.db.get_all_medical_reports()
        
        for report in reports:
            # Get test request to get patient and test type info
            test_request = self.db.get_test_request(report.test_request_id)
            patient_name = _("Unknown Patient")
            test_name = _("Unknown Test")
            status = _("Pending")
            
            if test_request:
                # Get patient info
                patient = self.db.get_patient(test_request.patient_id)
                if patient:
                    patient_name = patient.name
                
                # Get test type info
                test_type = self.db.get_test_type(test_request.test_type_id)
                if test_type:
                    test_name = test_type.name
                
                # Set status
                status = _(test_request.status.value)
            
            # Insert item and store the full ID in the item's values
            item_id = self.results_tree.insert("", tk.END, values=(
                report.id[:8],  # Short ID for display
                patient_name,
                test_name,
                status,
                report.created_at.strftime("%Y-%m-%d %H:%M")
            ))
            # Store the full ID in the item's tags for later retrieval
            self.results_tree.item(item_id, tags=(report.id,))
    
    def view_all_results(self):
        # This will show the reports screen
        self.show_reports()
    
    def create_new_result(self):
        # This will show the create report dialog
        self.create_report()
    
    def view_report(self):
        """View and edit a medical report"""
        selected = self.results_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a result to view"))
            return
        
        # Get the selected report ID from the item's tags
        item = self.results_tree.selection()[0]
        report_id = self.results_tree.item(item, "tags")[0]
        
        # Get report details from database
        report = self.db.get_medical_report(report_id)
        if not report:
            messagebox.showerror(_("Error"), _("Report not found"))
            return
        
        # Get test request info
        test_request = self.db.get_test_request(report.test_request_id)
        if not test_request:
            messagebox.showerror(_("Error"), _("Test request not found"))
            return
        
        # Get patient info
        patient = self.db.get_patient(test_request.patient_id)
        if not patient:
            messagebox.showerror(_("Error"), _("Patient not found"))
            return
        
        # Get test type info
        test_type = self.db.get_test_type(test_request.test_type_id)
        if not test_type:
            messagebox.showerror(_("Error"), _("Test type not found"))
            return
        
        # Create view/edit report dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("View/Edit Medical Report"))
        dialog.geometry("600x700")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Report header
        header_frame = ttk.Frame(dialog)
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("MEDICAL LABORATORY REPORT"), 
                 font=("Arial", 16, "bold")).pack()
        ttk.Label(header_frame, text=_("Report ID: {}").format(report.id), 
                 font=("Arial", 10)).pack()
        
        # Patient information
        patient_frame = ttk.LabelFrame(dialog, text=_("Patient Information"))
        patient_frame.pack(fill=tk.X, padx=10, pady=5)
        
        info_frame = ttk.Frame(patient_frame)
        info_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(info_frame, text=_("Patient: {}").format(patient.name)).grid(row=0, column=0, sticky=tk.W, padx=5)
        ttk.Label(info_frame, text=_("Age: {}").format(patient.age)).grid(row=0, column=1, sticky=tk.W, padx=5)
        ttk.Label(info_frame, text=_("Gender: {}").format(_(patient.gender.value))).grid(row=1, column=0, sticky=tk.W, padx=5)
        ttk.Label(info_frame, text=_("Contact: {}").format(patient.contact_info or _("N/A"))).grid(row=1, column=1, sticky=tk.W, padx=5)
        
        # Test information
        test_frame = ttk.LabelFrame(dialog, text=_("Test Information"))
        test_frame.pack(fill=tk.X, padx=10, pady=5)
        
        test_info_frame = ttk.Frame(test_frame)
        test_info_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(test_info_frame, text=_("Test: {}").format(test_type.name)).grid(row=0, column=0, sticky=tk.W, padx=5)
        ttk.Label(test_info_frame, text=_("Category: {}").format(test_type.category)).grid(row=0, column=1, sticky=tk.W, padx=5)
        ttk.Label(test_info_frame, text=_("Requested By: {}").format(test_request.requested_by)).grid(row=1, column=0, sticky=tk.W, padx=5)
        ttk.Label(test_info_frame, text=_("Requested At: {}").format(
            test_request.requested_at.strftime("%Y-%m-%d %H:%M"))).grid(row=1, column=1, sticky=tk.W, padx=5)
        
        # Result content with editing capability
        content_frame = ttk.LabelFrame(dialog, text=_("Result Details"))
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        content_text = tk.Text(content_frame, wrap=tk.WORD, height=15)
        content_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        content_text.insert("1.0", report.content)
        
        # Signature information
        signature_frame = ttk.Frame(dialog)
        signature_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(signature_frame, text=_("Signed By: {}").format(
            report.signed_by if report.signed_by != "N/A" else _("Not signed yet"))).pack(anchor=tk.W, padx=5)
        ttk.Label(signature_frame, text=_("Signed At: {}").format(
            report.signed_at.strftime("%Y-%m-%d %H:%M") if report.signed_at else _("Not signed yet"))).pack(anchor=tk.W, padx=5)
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        def save_changes():
            # Update report content
            new_content = content_text.get("1.0", tk.END).strip()
            report.content = new_content
            
            # Update signed info if not already signed
            if report.signed_by == "N/A":
                report.signed_by = self.current_user.id if self.current_user else "System"
                report.signed_at = datetime.now()
            
            # Save to database
            if self.db.update_medical_report(report):
                messagebox.showinfo(_("Success"), _("Report updated successfully"))
                dialog.destroy()
                self.load_results_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to update report"))
        
        def print_report():
            # In a real application, this would send the result to a printer
            # For now, we'll show a message
            messagebox.showinfo(_("Print"), _("In a real application, this would print the report.\n\n"
                                             "Patient: {}\n"
                                             "Test: {}\n"
                                             "Result: {}").format(
                                             patient.name, test_type.name, report.content[:50] + "..."))
        
        ttk.Button(button_frame, text=_("Save Changes"), command=save_changes, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Print"), command=print_report, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Close"), command=dialog.destroy, style="Accent.TButton").pack(side=tk.RIGHT, padx=5)
    
    def print_results(self):
        selected = self.results_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a result to print"))
            return
        
        # Get the selected report ID from the item's tags
        item = self.results_tree.selection()[0]
        report_id = self.results_tree.item(item, "tags")[0]
        
        # Get report details from database
        report = self.db.get_medical_report(report_id)
        
        if not report:
            messagebox.showerror(_("Error"), _("Report not found"))
            return
        
        # Get test request info
        test_request = self.db.get_test_request(report.test_request_id)
        if not test_request:
            messagebox.showerror(_("Error"), _("Test request not found"))
            return
        
        # Get patient info
        patient = self.db.get_patient(test_request.patient_id)
        if not patient:
            messagebox.showerror(_("Error"), _("Patient not found"))
            return
        
        # Get test type info
        test_type = self.db.get_test_type(test_request.test_type_id)
        if not test_type:
            messagebox.showerror(_("Error"), _("Test type not found"))
            return
        
        # Create print preview dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Print Result"))
        dialog.geometry("600x700")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Result header
        header_frame = ttk.Frame(dialog)
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("MEDICAL LABORATORY REPORT"), 
                 font=("Arial", 16, "bold")).pack()
        ttk.Label(header_frame, text=_("Result ID: {}").format(report.id), 
                 font=("Arial", 10)).pack()
        
        # Patient information
        patient_frame = ttk.LabelFrame(dialog, text=_("Patient Information"))
        patient_frame.pack(fill=tk.X, padx=10, pady=5)
        
        info_frame = ttk.Frame(patient_frame)
        info_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(info_frame, text=_("Patient: {}").format(patient.name)).grid(row=0, column=0, sticky=tk.W, padx=5)
        ttk.Label(info_frame, text=_("Age: {}").format(patient.age)).grid(row=0, column=1, sticky=tk.W, padx=5)
        ttk.Label(info_frame, text=_("Gender: {}").format(_(patient.gender.value))).grid(row=1, column=0, sticky=tk.W, padx=5)
        ttk.Label(info_frame, text=_("Contact: {}").format(patient.contact_info or _("N/A"))).grid(row=1, column=1, sticky=tk.W, padx=5)
        
        # Test information
        test_frame = ttk.LabelFrame(dialog, text=_("Test Information"))
        test_frame.pack(fill=tk.X, padx=10, pady=5)
        
        test_info_frame = ttk.Frame(test_frame)
        test_info_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(test_info_frame, text=_("Test: {}").format(test_type.name)).grid(row=0, column=0, sticky=tk.W, padx=5)
        ttk.Label(test_info_frame, text=_("Category: {}").format(test_type.category)).grid(row=0, column=1, sticky=tk.W, padx=5)
        ttk.Label(test_info_frame, text=_("Requested By: {}").format(test_request.requested_by)).grid(row=1, column=0, sticky=tk.W, padx=5)
        ttk.Label(test_info_frame, text=_("Requested At: {}").format(
            test_request.requested_at.strftime("%Y-%m-%d %H:%M"))).grid(row=1, column=1, sticky=tk.W, padx=5)
        
        # Result content
        content_frame = ttk.LabelFrame(dialog, text=_("Result Details"))
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        content_text = tk.Text(content_frame, wrap=tk.WORD)
        content_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        content_text.insert("1.0", report.content)
        
        # Signature information
        signature_frame = ttk.Frame(dialog)
        signature_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(signature_frame, text=_("Signed By: {}").format(
            report.signed_by if report.signed_by != "N/A" else _("Not signed yet"))).pack(anchor=tk.W, padx=5)
        ttk.Label(signature_frame, text=_("Signed At: {}").format(
            report.signed_at.strftime("%Y-%m-%d %H:%M") if report.signed_at else _("Not signed yet"))).pack(anchor=tk.W, padx=5)
        
        # Print buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(button_frame, text=_("Print"), 
                  command=lambda: self.do_print_result(dialog, report, patient, test_type, test_request)).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Close"), 
                  command=dialog.destroy).pack(side=tk.RIGHT, padx=5)

        header_frame = ttk.Frame(dialog)
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("MEDICAL LABORATORY REPORT"), 
                 font=("Arial", 16, "bold")).pack()
        ttk.Label(header_frame, text=_("Result ID: {}").format(report.id), 
                 font=("Arial", 10)).pack()
        
        # Patient information
        patient_frame = ttk.LabelFrame(dialog, text=_("Patient Information"))
        patient_frame.pack(fill=tk.X, padx=10, pady=5)
        
        info_frame = ttk.Frame(patient_frame)
        info_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(info_frame, text=_("Patient: {}").format(patient.name)).grid(row=0, column=0, sticky=tk.W, padx=5)
        ttk.Label(info_frame, text=_("Age: {}").format(patient.age)).grid(row=0, column=1, sticky=tk.W, padx=5)
        ttk.Label(info_frame, text=_("Gender: {}").format(_(patient.gender.value))).grid(row=1, column=0, sticky=tk.W, padx=5)
        ttk.Label(info_frame, text=_("Contact: {}").format(patient.contact_info or _("N/A"))).grid(row=1, column=1, sticky=tk.W, padx=5)
        
        # Test information
        test_frame = ttk.LabelFrame(dialog, text=_("Test Information"))
        test_frame.pack(fill=tk.X, padx=10, pady=5)
        
        test_info_frame = ttk.Frame(test_frame)
        test_info_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(test_info_frame, text=_("Test: {}").format(test_type.name)).grid(row=0, column=0, sticky=tk.W, padx=5)
        ttk.Label(test_info_frame, text=_("Category: {}").format(test_type.category)).grid(row=0, column=1, sticky=tk.W, padx=5)
        ttk.Label(test_info_frame, text=_("Requested By: {}").format(test_request.requested_by)).grid(row=1, column=0, sticky=tk.W, padx=5)
        ttk.Label(test_info_frame, text=_("Requested At: {}").format(
            test_request.requested_at.strftime("%Y-%m-%d %H:%M"))).grid(row=1, column=1, sticky=tk.W, padx=5)
        
        # Result content
        content_frame = ttk.LabelFrame(dialog, text=_("Result Details"))
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        content_text = tk.Text(content_frame, wrap=tk.WORD)
        content_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        content_text.insert("1.0", report.content)
        
        # Signature information
        signature_frame = ttk.Frame(dialog)
        signature_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(signature_frame, text=_("Signed By: {}").format(
            report.signed_by if report.signed_by != "N/A" else _("Not signed yet"))).pack(anchor=tk.W, padx=5)
        ttk.Label(signature_frame, text=_("Signed At: {}").format(
            report.signed_at.strftime("%Y-%m-%d %H:%M") if report.signed_at else _("Not signed yet"))).pack(anchor=tk.W, padx=5)
        
        # Print buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(button_frame, text=_("Print"), 
                  command=lambda: self.do_print_result(dialog, report, patient, test_type, test_request)).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Close"), 
                  command=dialog.destroy).pack(side=tk.RIGHT, padx=5)

    def remove_test(self):
        selected_index = self.test_list.curselection()
        if selected_index:
            self.test_list.delete(selected_index)

    def do_print_result(self, dialog, report, patient, test_type, test_request):
        # In a real application, this would send the result to a printer
        # For now, we'll show a message
        messagebox.showinfo(_("Print"), _("In a real application, this would print the result.\n\n"
                                         "Patient: {}\n"
                                         "Test: {}\n"
                                         "Result: {}").format(
                                         patient.name, test_type.name, report.content[:50] + "..."))
        dialog.destroy()
    
    def show_patients(self):
        self.current_screen = self.show_patients
        self.clear_content()
        
        # Patients header with 3D styling
        header_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("Patients Management"), 
                 style="Title.TLabel").pack(side=tk.LEFT)
        
        ttk.Button(header_frame, text=_("Add Patient"), 
                  command=self.add_patient, style="Accent.TButton").pack(side=tk.RIGHT)
        
        # Action buttons with 3D styling
        action_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        action_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Button(action_frame, text=_("View Details"), 
                  command=self.view_patient_details, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Edit Patient"), 
                  command=self.edit_patient, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Delete Patient"), 
                  command=self.delete_patient, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Request Test"), 
                  command=self.request_test_for_patient, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("View Test Requests"), 
                  command=self.view_patient_test_requests, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        
        # Patients table with enhanced styling
        table_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        table_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create treeview with custom styling
        columns = (_("ID"), _("Name"), _("Age"), _("Gender"), _("Contact"))
        self.patients_tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        
        for col in columns:
            self.patients_tree.heading(col, text=col)
            self.patients_tree.column(col, width=100)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, 
                                 command=self.patients_tree.yview)
        self.patients_tree.configure(yscroll=scrollbar.set)
        
        self.patients_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load patients data
        self.load_patients_data()
    
    def load_patients_data(self):
        # Check if patients_tree exists
        if not hasattr(self, 'patients_tree'):
            return
            
        # Clear existing data
        for item in self.patients_tree.get_children():
            self.patients_tree.delete(item)
        
        # Load patients from database
        patients = self.db.get_all_patients()
        
        for patient in patients:
            # Insert item with full 8-digit ID
            item_id = self.patients_tree.insert("", tk.END, values=(
                patient.id,  # Full 8-digit ID
                patient.name,
                patient.age,
                _(patient.gender.value),
                patient.contact_info
            ))
            # Store the full ID in the item's tags for later retrieval
            self.patients_tree.item(item_id, tags=(patient.id,))
    
    def add_patient(self):
        # Create add patient dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Add New Patient"))
        dialog.geometry("400x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Form fields
        ttk.Label(dialog, text=_("Name:")).pack(pady=5)
        name_entry = ttk.Entry(dialog, width=40)
        name_entry.pack(pady=5)
        
        ttk.Label(dialog, text=_("Age:")).pack(pady=5)
        age_entry = ttk.Entry(dialog, width=40)
        age_entry.pack(pady=5)
        
        ttk.Label(dialog, text=_("Gender:")).pack(pady=5)
        gender_var = tk.StringVar()
        gender_combo = ttk.Combobox(dialog, textvariable=gender_var,
                                   values=[_("Male"), _("Female"), _("Other")],
                                   state="readonly", width=37)
        gender_combo.pack(pady=5)
        
        ttk.Label(dialog, text=_("Contact Info:")).pack(pady=5)
        contact_entry = ttk.Entry(dialog, width=40)
        contact_entry.pack(pady=5)
        
        def save_patient():
            name = name_entry.get().strip()
            age_str = age_entry.get().strip()
            gender_text = gender_var.get()
            contact = contact_entry.get().strip()
            
            # Validation
            if not name:
                messagebox.showerror(_("Error"), _("Please enter patient name"))
                return
            
            try:
                age = int(age_str)
                if age < 0 or age > 150:
                    raise ValueError()
            except ValueError:
                messagebox.showerror(_("Error"), _("Please enter a valid age"))
                return
            
            if not gender_text:
                messagebox.showerror(_("Error"), _("Please select gender"))
                return
            
            # Map gender text to enum
            gender_map = {
                _("Male"): Gender.MALE,
                _("Female"): Gender.FEMALE,
                _("Other"): Gender.OTHER
            }
            gender = gender_map.get(gender_text, Gender.OTHER)
            
            # Generate 8-digit patient ID
            patient_id = self.db.generate_patient_id()
            
            # Create patient
            patient = Patient(
                id=patient_id,
                name=name,
                age=age,
                gender=gender,
                contact_info=contact
            )
            
            result = self.db.create_patient(patient)
            if result:
                messagebox.showinfo(_("Success"), _("Patient added successfully with ID: {}").format(patient_id))
                dialog.destroy()
                # Force refresh of patient data
                self.load_patients_data()  # Refresh only the data, not the entire view
            else:
                messagebox.showerror(_("Error"), _("Failed to add patient"))
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text=_("Save"), command=save_patient).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), 
                  command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Focus on first entry
        name_entry.focus()
    
    def view_patient_details(self):
        selected = self.patients_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a patient"))
            return
        
        # Get the selected patient ID from the item's tags
        item = self.patients_tree.selection()[0]
        patient_id = self.patients_tree.item(item, "tags")[0]
        
        # Get patient details from database
        patient = self.db.get_patient(patient_id)
        if not patient:
            messagebox.showerror(_("Error"), _("Patient not found"))
            return
        
        # Create patient details dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Patient Details"))
        dialog.geometry("400x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Display patient information
        ttk.Label(dialog, text=_("Patient ID:"), font=("Arial", 10, "bold")).pack(anchor=tk.W, padx=10, pady=(10, 0))
        ttk.Label(dialog, text=patient.id).pack(anchor=tk.W, padx=20)
        
        ttk.Label(dialog, text=_("Name:"), font=("Arial", 10, "bold")).pack(anchor=tk.W, padx=10, pady=(10, 0))
        ttk.Label(dialog, text=patient.name).pack(anchor=tk.W, padx=20)
        
        ttk.Label(dialog, text=_("Age:"), font=("Arial", 10, "bold")).pack(anchor=tk.W, padx=10, pady=(10, 0))
        ttk.Label(dialog, text=str(patient.age)).pack(anchor=tk.W, padx=20)
        
        ttk.Label(dialog, text=_("Gender:"), font=("Arial", 10, "bold")).pack(anchor=tk.W, padx=10, pady=(10, 0))
        ttk.Label(dialog, text=_(patient.gender.value)).pack(anchor=tk.W, padx=20)
        
        ttk.Label(dialog, text=_("Contact Info:"), font=("Arial", 10, "bold")).pack(anchor=tk.W, padx=10, pady=(10, 0))
        ttk.Label(dialog, text=patient.contact_info or _("N/A")).pack(anchor=tk.W, padx=20)
        
        ttk.Button(dialog, text=_("Close"), command=dialog.destroy).pack(pady=20)
    
    def edit_patient(self):
        selected = self.patients_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a patient"))
            return
        
        # Get the selected patient ID from the item's tags
        item = self.patients_tree.selection()[0]
        patient_id = self.patients_tree.item(item, "tags")[0]
        
        # Get patient details from database
        patient = self.db.get_patient(patient_id)
        if not patient:
            messagebox.showerror(_("Error"), _("Patient not found"))
            return
        
        # Create edit patient dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Edit Patient"))
        dialog.geometry("400x350")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Form fields with current patient data
        ttk.Label(dialog, text=_("Name:")).pack(pady=5)
        name_entry = ttk.Entry(dialog, width=40)
        name_entry.pack(pady=5)
        name_entry.insert(0, patient.name)
        
        ttk.Label(dialog, text=_("Age:")).pack(pady=5)
        age_entry = ttk.Entry(dialog, width=40)
        age_entry.pack(pady=5)
        age_entry.insert(0, str(patient.age))
        
        ttk.Label(dialog, text=_("Gender:")).pack(pady=5)
        gender_var = tk.StringVar(value=_(patient.gender.value))
        gender_combo = ttk.Combobox(dialog, textvariable=gender_var,
                                   values=[_("Male"), _("Female"), _("Other")],
                                   state="readonly", width=37)
        gender_combo.pack(pady=5)
        
        ttk.Label(dialog, text=_("Contact Info:")).pack(pady=5)
        contact_entry = ttk.Entry(dialog, width=40)
        contact_entry.pack(pady=5)
        contact_entry.insert(0, patient.contact_info or "")
        
        def save_patient():
            name = name_entry.get().strip()
            age_str = age_entry.get().strip()
            gender_text = gender_var.get()
            contact = contact_entry.get().strip()
            
            # Validation
            if not name:
                messagebox.showerror(_("Error"), _("Please enter patient name"))
                return
            
            try:
                age = int(age_str)
                if age < 0 or age > 150:
                    raise ValueError()
            except ValueError:
                messagebox.showerror(_("Error"), _("Please enter a valid age (0-150)"))
                return
            
            if not gender_text:
                messagebox.showerror(_("Error"), _("Please select gender"))
                return
            
            # Map gender text to enum
            gender_map = {
                _("Male"): Gender.MALE,
                _("Female"): Gender.FEMALE,
                _("Other"): Gender.OTHER
            }
            gender = gender_map.get(gender_text, Gender.OTHER)
            
            # Update patient
            patient.name = name
            patient.age = age
            patient.gender = gender
            patient.contact_info = contact
            
            if self.db.update_patient(patient):
                messagebox.showinfo(_("Success"), _("Patient updated successfully"))
                dialog.destroy()
                self.load_patients_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to update patient"))
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text=_("Save"), command=save_patient).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), 
                  command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Focus on first entry
        name_entry.focus()
    
    def delete_patient(self):
        selected = self.patients_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a patient"))
            return
        
        # Get the selected patient ID from the item's tags
        item = self.patients_tree.selection()[0]
        patient_id = self.patients_tree.item(item, "tags")[0]
        
        # Confirm deletion
        if messagebox.askyesno(_("Confirm Delete"), 
                              _("Are you sure you want to delete this patient?")):
            if self.db.delete_patient(patient_id):
                messagebox.showinfo(_("Success"), _("Patient deleted successfully"))
                self.load_patients_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to delete patient"))
    
    def request_test_for_patient(self):
        selected = self.patients_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a patient"))
            return
        
        # Get the selected patient ID from the item's tags
        item = self.patients_tree.selection()[0]
        patient_id = self.patients_tree.item(item, "tags")[0]
        
        # Get patient details
        patient = self.db.get_patient(patient_id)
        if not patient:
            messagebox.showerror(_("Error"), _("Patient not found"))
            return
        
        # Create professional test request dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("🏥 Request Medical Examinations for Patient"))
        dialog.geometry("600x550")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.configure(bg="#f5f7fa")
        
        # Professional header
        header_frame = ttk.Frame(dialog, style="Header.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("📋 Medical Examination Request"), 
                 font=("Arial", 16, "bold"), style="Header.TLabel").pack(pady=10)
        
        # Patient information card
        patient_frame = ttk.LabelFrame(dialog, text=_("👤 Patient Information"), padding=15)
        patient_frame.pack(fill=tk.X, padx=15, pady=5)
        
        patient_info = ttk.Frame(patient_frame)
        patient_info.pack(fill=tk.X)
        
        ttk.Label(patient_info, text=_("Patient Name:"), font=("Arial", 10, "bold")).grid(row=0, column=0, sticky=tk.W, padx=5, pady=2)
        ttk.Label(patient_info, text=patient.name, font=("Arial", 10)).grid(row=0, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(patient_info, text=_("Patient ID:"), font=("Arial", 10, "bold")).grid(row=0, column=2, sticky=tk.W, padx=5, pady=2)
        ttk.Label(patient_info, text=patient.id, font=("Arial", 10)).grid(row=0, column=3, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(patient_info, text=_("Age:"), font=("Arial", 10, "bold")).grid(row=1, column=0, sticky=tk.W, padx=5, pady=2)
        ttk.Label(patient_info, text=str(patient.age), font=("Arial", 10)).grid(row=1, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(patient_info, text=_("Gender:"), font=("Arial", 10, "bold")).grid(row=1, column=2, sticky=tk.W, padx=5, pady=2)
        ttk.Label(patient_info, text=_(patient.gender.value), font=("Arial", 10)).grid(row=1, column=3, sticky=tk.W, padx=5, pady=2)
        
        # Test selection section with multiple selection capability
        test_frame = ttk.LabelFrame(dialog, text=_("🔍 Select Medical Examinations"), padding=15)
        test_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)
        
        # Available tests listbox with multiple selection
        ttk.Label(test_frame, text=_("Available Examinations:"), font=("Arial", 11, "bold")).pack(anchor=tk.W, pady=(0, 5))
        
        # Create a frame for the listbox and scrollbar
        listbox_frame = ttk.Frame(test_frame)
        listbox_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Listbox with multiple selection
        available_listbox = tk.Listbox(listbox_frame, selectmode=tk.MULTIPLE, height=10, 
                                      font=("Arial", 10), bg="white", relief="solid", bd=1)
        available_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        
        # Scrollbar
        listbox_scrollbar = ttk.Scrollbar(listbox_frame, orient=tk.VERTICAL, command=available_listbox.yview)
        available_listbox.configure(yscrollcommand=listbox_scrollbar.set)
        listbox_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Selection controls
        selection_frame = ttk.Frame(test_frame)
        selection_frame.pack(fill=tk.X, pady=10)
        
        def select_all():
            available_listbox.selection_set(0, tk.END)
            update_selected_count()
        
        def deselect_all():
            available_listbox.selection_clear(0, tk.END)
            update_selected_count()
        
        ttk.Button(selection_frame, text=_("✓ Select All"), command=select_all, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(selection_frame, text=_("✗ Deselect All"), command=deselect_all, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        
        # Selected count indicator
        selected_count_frame = ttk.Frame(test_frame)
        selected_count_frame.pack(fill=tk.X, pady=5)
        
        selected_count_label = ttk.Label(selected_count_frame, 
                                       text=_("Selected examinations: 0"), 
                                       font=("Arial", 10, "bold"), foreground="#3498db")
        selected_count_label.pack(side=tk.LEFT)
        
        # Load available tests from database
        test_types = self.db.get_all_test_types()
        test_map = {}  # Map display text to test objects
        
        if not test_types:
            messagebox.showwarning(_("Warning"), _("No examination types available. Please add examination types first."))
            dialog.destroy()
            return
        
        for test in test_types:
            display_text = f"🩺 {test.name} - {test.category} (${test.price:.2f})"
            available_listbox.insert(tk.END, display_text)
            test_map[display_text] = test
        
        def update_selected_count():
            selected_count = len(available_listbox.curselection())
            selected_count_label.config(text=_("Selected examinations: {}").format(selected_count))
        
        # Bind selection event to update count
        available_listbox.bind('<<ListboxSelect>>', lambda e: update_selected_count())
        
        # Request details section
        details_frame = ttk.LabelFrame(dialog, text=_("📝 Request Details"), padding=15)
        details_frame.pack(fill=tk.X, padx=15, pady=5)
        
        # Requested by
        ttk.Label(details_frame, text=_("Requested By:"), font=("Arial", 10, "bold")).grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        requested_by_entry = ttk.Entry(details_frame, width=30, font=("Arial", 10))
        requested_by_entry.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        requested_by_entry.insert(0, self.current_user.username if self.current_user else "")
        
        # Request date (default to current date)
        ttk.Label(details_frame, text=_("Request Date:"), font=("Arial", 10, "bold")).grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        request_date_entry = ttk.Entry(details_frame, width=30, font=("Arial", 10))
        request_date_entry.grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)
        request_date_entry.insert(0, datetime.now().strftime("%Y-%m-%d %H:%M"))
        request_date_entry.config(state="readonly")  # Make it readonly as it's auto-generated
        
        # Additional notes
        ttk.Label(details_frame, text=_("Notes:"), font=("Arial", 10, "bold")).grid(row=2, column=0, sticky=tk.NW, padx=5, pady=5)
        notes_text = tk.Text(details_frame, width=30, height=3, font=("Arial", 10), wrap=tk.WORD)
        notes_text.grid(row=2, column=1, sticky=tk.W, padx=5, pady=5)
        
        def save_test_requests():
            # Get selected tests
            selected_indices = available_listbox.curselection()
            if not selected_indices:
                messagebox.showerror(_("Error"), _("Please select at least one examination"))
                return
            
            requested_by = requested_by_entry.get().strip()
            if not requested_by:
                messagebox.showerror(_("Error"), _("Please enter who requested the examinations"))
                return
            
            # Get notes
            notes = notes_text.get("1.0", tk.END).strip()
            
            # Create test requests for each selected test
            success_count = 0
            failed_tests = []
            
            for index in selected_indices:
                display_text = available_listbox.get(index)
                test_type = test_map.get(display_text)
                
                if test_type:
                    # Create test request
                    test_request = TestRequest(
                        id=str(uuid.uuid4()),
                        patient_id=patient_id,
                        test_type_id=test_type.id,
                        requested_by=requested_by,
                        requested_at=datetime.now(),
                        status=TestStatus.PENDING
                    )
                    
                    if self.db.create_test_request(test_request):
                        success_count += 1
                        # If there are notes, we could save them as part of the request
                        # This would require modifying the database schema
                    else:
                        failed_tests.append(test_type.name)
            
            if success_count > 0:
                if failed_tests:
                    messagebox.showwarning(_("Partial Success"), 
                                         _("{} examination(s) requested successfully. Failed to request: {}").format(
                                         success_count, ", ".join(failed_tests)))
                else:
                    messagebox.showinfo(_("Success"), 
                                      _("{} examination(s) requested successfully").format(success_count))
                dialog.destroy()
                # Refresh the UI if needed
                if hasattr(self, 'load_patients_data'):
                    self.load_patients_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to request examinations: {}").format(", ".join(failed_tests)))
        
        # Action buttons with professional styling
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=15, pady=15)
        
        ttk.Button(button_frame, text=_("💾 Save Request"), command=save_test_requests, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("❌ Cancel"), 
                  command=dialog.destroy, style="Accent.TButton").pack(side=tk.RIGHT, padx=5)
    
    def view_patient_test_requests(self):
        selected = self.patients_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a patient"))
            return
        
        # Get the selected patient ID from the item's tags
        item = self.patients_tree.selection()[0]
        patient_id = self.patients_tree.item(item, "tags")[0]
        
        # Get patient details
        patient = self.db.get_patient(patient_id)
        if not patient:
            messagebox.showerror(_("Error"), _("Patient not found"))
            return
        
        # Create professional test requests dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("📋 Test Requests for Patient: {}").format(patient.name))
        dialog.geometry("700x500")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.configure(bg="#f5f7fa")
        
        # Professional header
        header_frame = ttk.Frame(dialog, style="Header.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("📋 Medical Test Requests"), 
                 font=("Arial", 16, "bold"), style="Header.TLabel").pack(pady=10)
        
        # Patient information card
        patient_frame = ttk.LabelFrame(dialog, text=_("👤 Patient Information"), padding=15)
        patient_frame.pack(fill=tk.X, padx=15, pady=5)
        
        patient_info = ttk.Frame(patient_frame)
        patient_info.pack(fill=tk.X)
        
        ttk.Label(patient_info, text=_("Patient Name:"), font=("Arial", 10, "bold")).grid(row=0, column=0, sticky=tk.W, padx=5, pady=2)
        ttk.Label(patient_info, text=patient.name, font=("Arial", 10)).grid(row=0, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(patient_info, text=_("Patient ID:"), font=("Arial", 10, "bold")).grid(row=0, column=2, sticky=tk.W, padx=5, pady=2)
        ttk.Label(patient_info, text=patient.id, font=("Arial", 10)).grid(row=0, column=3, sticky=tk.W, padx=5, pady=2)
        
        # Test requests table with enhanced styling
        table_frame = ttk.LabelFrame(dialog, text=_("📋 Test Requests"), padding=10)
        table_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)
        
        # Create treeview with enhanced styling
        columns = (_("Request ID"), _("Test Type"), _("Requested By"), _("Requested At"), _("Status"))
        requests_tree = ttk.Treeview(table_frame, columns=columns, show="headings", height=12)
        
        # Configure column headings and widths
        requests_tree.heading(_("Request ID"), text=_("Request ID"))
        requests_tree.column(_("Request ID"), width=100, anchor=tk.CENTER)
        requests_tree.heading(_("Test Type"), text=_("Test Type"))
        requests_tree.column(_("Test Type"), width=150, anchor=tk.W)
        requests_tree.heading(_("Requested By"), text=_("Requested By"))
        requests_tree.column(_("Requested By"), width=120, anchor=tk.W)
        requests_tree.heading(_("Requested At"), text=_("Requested At"))
        requests_tree.column(_("Requested At"), width=120, anchor=tk.CENTER)
        requests_tree.heading(_("Status"), text=_("Status"))
        requests_tree.column(_("Status"), width=100, anchor=tk.CENTER)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=requests_tree.yview)
        requests_tree.configure(yscroll=scrollbar.set)
        
        requests_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load test requests for this patient
        test_requests = self.db.get_test_requests_by_patient(patient_id)
        request_map = {}  # Map item IDs to request objects
        
        for request in test_requests:
            # Get test type name
            test_type = self.db.get_test_type(request.test_type_id)
            test_name = test_type.name if test_type else _("Unknown Test")
            
            # Insert item and store the full ID in the item's values
            item_id = requests_tree.insert("", tk.END, values=(
                request.id[:8],
                test_name,
                request.requested_by,
                request.requested_at.strftime("%Y-%m-%d %H:%M"),
                _(request.status.value)
            ))
            # Store the full request object in our map
            request_map[item_id] = request
        
        # Store references for button functions
        dialog.requests_tree = requests_tree
        dialog.request_map = request_map
        dialog.patient_id = patient_id
        
        def edit_selected_request():
            selected = requests_tree.selection()
            if not selected:
                messagebox.showwarning(_("Warning"), _("Please select a test request to edit"))
                return
            
            item_id = selected[0]
            request = request_map.get(item_id)
            if not request:
                messagebox.showerror(_("Error"), _("Test request not found"))
                return
            
            # Open edit dialog
            self.edit_test_request(request, dialog)
        
        def delete_selected_request():
            selected = requests_tree.selection()
            if not selected:
                messagebox.showwarning(_("Warning"), _("Please select a test request to delete"))
                return
            
            item_id = selected[0]
            request = request_map.get(item_id)
            if not request:
                messagebox.showerror(_("Error"), _("Test request not found"))
                return
            
            # Confirm deletion
            test_type = self.db.get_test_type(request.test_type_id)
            test_name = test_type.name if test_type else _("Unknown Test")
            
            result = messagebox.askyesno(
                _("Confirm Deletion"),
                _("Are you sure you want to delete the test request for '{}'?").format(test_name)
            )
            
            if result:
                if self.db.delete_test_request(request.id):
                    messagebox.showinfo(_("Success"), _("Test request deleted successfully"))
                    # Remove from tree and map
                    requests_tree.delete(item_id)
                    del request_map[item_id]
                else:
                    messagebox.showerror(_("Error"), _("Failed to delete test request"))
        
        # Action buttons with professional styling
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=15, pady=10)
        
        ttk.Button(button_frame, text=_("✏️ Edit Request"), command=edit_selected_request, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("🗑️ Delete Request"), command=delete_selected_request, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("❌ Close"), command=dialog.destroy, 
                  style="Accent.TButton").pack(side=tk.RIGHT, padx=5)
    
    def edit_test_request(self, request, parent_dialog):
        """Edit a test request"""
        # Get test type
        test_type = self.db.get_test_type(request.test_type_id)
        if not test_type:
            messagebox.showerror(_("Error"), _("Test type not found"))
            return
        
        # Create edit dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("✏️ Edit Test Request"))
        dialog.geometry("450x350")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.configure(bg="#f5f7fa")
        
        # Professional header
        header_frame = ttk.Frame(dialog, style="Header.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("Edit Test Request"), 
                 font=("Arial", 14, "bold"), style="Header.TLabel").pack(pady=5)
        
        # Test information
        info_frame = ttk.LabelFrame(dialog, text=_("📋 Test Information"), padding=15)
        info_frame.pack(fill=tk.X, padx=15, pady=5)
        
        ttk.Label(info_frame, text=_("Test Type:"), font=("Arial", 10, "bold")).grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Label(info_frame, text=test_type.name, font=("Arial", 10)).grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Status selection
        ttk.Label(info_frame, text=_("Status:"), font=("Arial", 10, "bold")).grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        status_var = tk.StringVar(value=request.status.value)
        status_combo = ttk.Combobox(info_frame, textvariable=status_var,
                                   values=[_("Pending"), _("In Progress"), _("Completed"), _("Cancelled")],
                                   state="readonly", width=25)
        status_combo.grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Requested by
        ttk.Label(info_frame, text=_("Requested By:"), font=("Arial", 10, "bold")).grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        requested_by_entry = ttk.Entry(info_frame, width=28, font=("Arial", 10))
        requested_by_entry.grid(row=2, column=1, sticky=tk.W, padx=5, pady=5)
        requested_by_entry.insert(0, request.requested_by)
        
        # Request date (readonly)
        ttk.Label(info_frame, text=_("Request Date:"), font=("Arial", 10, "bold")).grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
        request_date_entry = ttk.Entry(info_frame, width=28, font=("Arial", 10))
        request_date_entry.grid(row=3, column=1, sticky=tk.W, padx=5, pady=5)
        request_date_entry.insert(0, request.requested_at.strftime("%Y-%m-%d %H:%M"))
        request_date_entry.config(state="readonly")
        
        def save_changes():
            # Get values
            requested_by = requested_by_entry.get().strip()
            if not requested_by:
                messagebox.showerror(_("Error"), _("Please enter who requested the test"))
                return
            
            # Map status text to enum
            status_map = {
                _("Pending"): TestStatus.PENDING,
                _("In Progress"): TestStatus.IN_PROGRESS,
                _("Completed"): TestStatus.COMPLETED,
                _("Cancelled"): TestStatus.CANCELLED
            }
            status = status_map.get(status_var.get(), TestStatus.PENDING)
            
            # Update request object
            request.requested_by = requested_by
            request.status = status
            
            # Update in database
            if self.db.update_test_request(request):
                messagebox.showinfo(_("Success"), _("Test request updated successfully"))
                dialog.destroy()
                
                # Update the parent dialog
                for item_id, req in parent_dialog.request_map.items():
                    if req.id == request.id:
                        # Update the tree view
                        test_name = test_type.name if test_type else _("Unknown Test")
                        parent_dialog.requests_tree.item(item_id, values=(
                            request.id[:8],
                            test_name,
                            request.requested_by,
                            request.requested_at.strftime("%Y-%m-%d %H:%M"),
                            _(request.status.value)
                        ))
                        # Update the map
                        parent_dialog.request_map[item_id] = request
                        break
            else:
                messagebox.showerror(_("Error"), _("Failed to update test request"))
        
        # Action buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=15, pady=15)
        
        ttk.Button(button_frame, text=_("💾 Save Changes"), command=save_changes, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("❌ Cancel"), command=dialog.destroy, 
                  style="Accent.TButton").pack(side=tk.RIGHT, padx=5)
    
    def edit_system_name(self):
        """Allow admin to edit the system name"""
        if not self.current_user or self.current_user.role != UserRole.ADMIN:
            messagebox.showerror(_("Error"), _("Only administrators can modify the system name"))
            return
        
        # Create dialog to edit system name
        dialog = tk.Toplevel(self.root)
        dialog.title(_("✏️ Edit System Name"))
        dialog.geometry("400x150")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.configure(bg="#f5f7fa")
        
        # Professional header
        header_frame = ttk.Frame(dialog, style="Header.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("Edit System Name"), 
                 font=("Arial", 14, "bold"), style="Header.TLabel").pack(pady=5)
        
        # Current name
        ttk.Label(dialog, text=_("Current System Name:"), font=("Arial", 10, "bold")).pack(pady=(10, 0))
        current_name_label = ttk.Label(dialog, text=self.title_label.cget("text"), font=("Arial", 10))
        current_name_label.pack()
        
        # New name entry
        ttk.Label(dialog, text=_("New System Name:"), font=("Arial", 10, "bold")).pack(pady=(10, 0))
        new_name_entry = ttk.Entry(dialog, width=40, font=("Arial", 10))
        new_name_entry.pack(pady=5)
        new_name_entry.insert(0, self.title_label.cget("text"))
        new_name_entry.focus()
        
        def save_system_name():
            new_name = new_name_entry.get().strip()
            if not new_name:
                messagebox.showerror(_("Error"), _("Please enter a system name"))
                return
            
            # Update the system name in the UI
            self.title_label.config(text=new_name)
            self.root.title(new_name)
            
            # In a real application, you might want to save this to a configuration file or database
            messagebox.showinfo(_("Success"), _("System name updated successfully"))
            dialog.destroy()
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=15)
        
        ttk.Button(button_frame, text=_("Save"), command=save_system_name, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), command=dialog.destroy, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)

    def show_results(self):
        self.current_screen = self.show_results
        self.clear_content()
        
        # Results header with 3D styling
        header_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("Results Management"), 
                 style="Title.TLabel").pack(side=tk.LEFT)
        
        # Add button frame for multiple actions
        button_frame = ttk.Frame(header_frame)
        button_frame.pack(side=tk.RIGHT)
        
        ttk.Button(button_frame, text=_("Add Result"), 
                  command=self.create_new_result, style="Accent.TButton").pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text=_("Manage Templates"), 
                  command=self.manage_test_templates, style="Accent.TButton").pack(side=tk.LEFT, padx=2)
        
        # Results table with enhanced styling
        table_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        table_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create treeview with custom styling
        columns = (_("ID"), _("Patient"), _("Test Type"), _("Status"), _("Created At"))
        self.results_tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        
        for col in columns:
            self.results_tree.heading(col, text=col)
            self.results_tree.column(col, width=120)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, 
                                 command=self.results_tree.yview)
        self.results_tree.configure(yscroll=scrollbar.set)
        
        self.results_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load results data
        self.load_results_data()
        
        # Action buttons with 3D styling
        action_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        action_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(action_frame, text=_("View Details"), 
                  command=self.view_result_details, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Edit Result"), 
                  command=self.edit_result, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Delete Result"), 
                  command=self.delete_result, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Print Result"), 
                  command=self.print_selected_result, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
    
    def show_tests(self):
        self.current_screen = self.show_tests
        self.clear_content()
        
        # Tests header with 3D styling
        header_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("Tests Management"), 
                 style="Title.TLabel").pack(side=tk.LEFT)
        
        ttk.Button(header_frame, text=_("Add Test"), 
                  command=self.add_test, style="Accent.TButton").pack(side=tk.RIGHT)
        
        # Tests table with enhanced styling
        table_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        table_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create treeview with custom styling
        columns = (_("ID"), _("Name"), _("Category"), _("Price"), _("Description"))
        self.tests_tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        
        for col in columns:
            self.tests_tree.heading(col, text=col)
            self.tests_tree.column(col, width=120)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, 
                                 command=self.tests_tree.yview)
        self.tests_tree.configure(yscroll=scrollbar.set)
        
        self.tests_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load tests data
        self.load_tests_data()
        
        # Action buttons with 3D styling
        action_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        action_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(action_frame, text=_("View Details"), 
                  command=self.view_test_details, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Edit Test"), 
                  command=self.edit_test, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Delete Test"), 
                  command=self.delete_test, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
    
    def load_tests_data(self):
        # Check if tests_tree exists
        if not hasattr(self, 'tests_tree'):
            return
            
        # Clear existing data
        for item in self.tests_tree.get_children():
            self.tests_tree.delete(item)
        
        # Load test types from database
        test_types = self.db.get_all_test_types()
        
        for test in test_types:
            # Format the ID to ensure it's displayed as a three-digit number
            display_id = test.id if len(test.id) == 3 and test.id.isdigit() else test.id[:8]
            
            # Insert item and store the full ID in the item's tags for later retrieval
            item_id = self.tests_tree.insert("", tk.END, values=(
                display_id,  # Show three-digit ID or short ID
                test.name,
                _(test.category),
                f"${test.price:.2f}",
                test.description[:50] + "..." if len(test.description) > 50 else test.description
            ))
            # Store the full ID in the item's tags for later retrieval
            self.tests_tree.item(item_id, tags=(test.id,))
    
    def add_test(self):
        # Create add test dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Add New Test"))
        dialog.geometry("400x350")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Form fields
        ttk.Label(dialog, text=_("Test Name:")).pack(pady=5)
        name_entry = ttk.Entry(dialog, width=40)
        name_entry.pack(pady=5)
        
        ttk.Label(dialog, text=_("Category:")).pack(pady=5)
        category_var = tk.StringVar()
        category_combo = ttk.Combobox(dialog, textvariable=category_var,
                                      values=[_("Blood"), _("Urine"), _("Biochemistry"), _("Imaging"), _("Genetics"), _("Hematology"), _("Microbiology"), _("Others")], state="readonly", width=37)
        category_combo.pack(pady=5)
        
        ttk.Label(dialog, text=_("Price:")).pack(pady=5)
        price_entry = ttk.Entry(dialog, width=40)
        price_entry.pack(pady=5)
        
        ttk.Label(dialog, text=_("Description:")).pack(pady=5)
        description_entry = ttk.Entry(dialog, width=40)
        description_entry.pack(pady=5)
        
        def save_test():
            name = name_entry.get().strip()
            category = category_var.get().strip()
            price = price_entry.get().strip()
            description = description_entry.get().strip()
            
            # Validation
            if not name:
                messagebox.showerror(_("Error"), _("Please enter a test name"))
                return
            
            if not category:
                messagebox.showerror(_("Error"), _("Please select a category"))
                return
            
            if not price:
                messagebox.showerror(_("Error"), _("Please enter a price"))
                return
            
            try:
                price = float(price)
            except ValueError:
                messagebox.showerror(_("Error"), _("Price must be a number"))
                return
            
            if price < 0:
                messagebox.showerror(_("Error"), _("Price cannot be negative"))
                return
            
            # Create test type
            test_type = TestType(
                id=str(uuid.uuid4()),
                name=name,
                category=category,
                price=price,
                description=description
            )
            
            if self.db.create_test_type(test_type):
                messagebox.showinfo(_("Success"), _("Test added successfully"))
                dialog.destroy()
                self.load_tests_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to add test"))
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text=_("Save"), command=save_test).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), 
                  command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Focus on first entry
        name_entry.focus()
    
    def view_test_details(self):
        selected = self.tests_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a test"))
            return
        
        # Get the selected item
        item = self.tests_tree.selection()[0]
        
        # Try to get the test ID from tags first, fallback to values if needed
        try:
            test_id = self.tests_tree.item(item, "tags")[0]
        except (IndexError, TypeError):
            # If tags don't contain the ID, try to get it from the values
            values = self.tests_tree.item(item, "values")
            if values and len(values) > 0:
                test_id = values[0]  # ID is the first column
            else:
                messagebox.showerror(_("Error"), _("Unable to identify test"))
                return
        
        # Get test details
        test = self.db.get_test_type(test_id)
        if not test:
            # Try to find test by display ID
            all_tests = self.db.get_all_test_types()
            test = next((t for t in all_tests if t.id.startswith(test_id) or t.id[:8] == test_id), None)
            
            if not test:
                messagebox.showerror(_("Error"), _("Test not found"))
                return
        
        # Create test details dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Test Details: {}").format(test.name))
        dialog.geometry("400x350")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Test details
        ttk.Label(dialog, text=_("Test Name: {}").format(test.name)).pack(pady=5)
        ttk.Label(dialog, text=_("Category: {}").format(test.category)).pack(pady=5)
        ttk.Label(dialog, text=_("Price: ${:.2f}").format(test.price)).pack(pady=5)
        ttk.Label(dialog, text=_("Description: {}").format(test.description)).pack(pady=5)
        
        # Close button
        ttk.Button(dialog, text=_("Close"), command=dialog.destroy).pack(pady=10)

    def edit_test(self):
        selected = self.tests_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a test"))
            return
        
        # Get the selected item
        item = self.tests_tree.selection()[0]
        
        # Try to get the test ID from tags first, fallback to values if needed
        try:
            test_id = self.tests_tree.item(item, "tags")[0]
        except (IndexError, TypeError):
            # If tags don't contain the ID, try to get it from the values
            values = self.tests_tree.item(item, "values")
            if values and len(values) > 0:
                test_id = values[0]  # ID is the first column
            else:
                messagebox.showerror(_("Error"), _("Unable to identify test"))
                return
        
        # Get test details
        test = self.db.get_test_type(test_id)
        if not test:
            # Try to find test by display ID
            all_tests = self.db.get_all_test_types()
            test = next((t for t in all_tests if t.id.startswith(test_id) or t.id[:8] == test_id), None)
            
            if not test:
                messagebox.showerror(_("Error"), _("Test not found"))
                return
        
        # Create edit test dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Edit Test: {}").format(test.name))
        dialog.geometry("400x350")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Form fields
        ttk.Label(dialog, text=_("Test Name:")).pack(pady=5)
        name_entry = ttk.Entry(dialog, width=40)
        name_entry.pack(pady=5)
        name_entry.insert(0, test.name)
        
        ttk.Label(dialog, text=_("Category:")).pack(pady=5)
        category_var = tk.StringVar()
        category_combo = ttk.Combobox(dialog, textvariable=category_var,
                                      values=[_("Blood"), _("Urine"), _("Biochemistry"), _("Imaging"), _("Genetics"), _("Hematology"), _("Microbiology"), _("Others")], state="readonly", width=37)
        category_combo.pack(pady=5)
        category_combo.set(test.category)
        
        ttk.Label(dialog, text=_("Price:")).pack(pady=5)
        price_entry = ttk.Entry(dialog, width=40)
        price_entry.pack(pady=5)
        price_entry.insert(0, str(test.price))
        
        ttk.Label(dialog, text=_("Description:")).pack(pady=5)
        description_entry = ttk.Entry(dialog, width=40)
        description_entry.pack(pady=5)
        description_entry.insert(0, test.description)
        
        def save_test():
            name = name_entry.get().strip()
            category = category_var.get().strip()
            price = price_entry.get().strip()
            description = description_entry.get().strip()
            
            # Validation
            if not name:
                messagebox.showerror(_("Error"), _("Please enter a test name"))
                return
            
            if not category:
                messagebox.showerror(_("Error"), _("Please select a category"))
                return
            
            if not price:
                messagebox.showerror(_("Error"), _("Please enter a price"))
                return
            
            try:
                price = float(price)
            except ValueError:
                messagebox.showerror(_("Error"), _("Price must be a number"))
                return
            
            if price < 0:
                messagebox.showerror(_("Error"), _("Price cannot be negative"))
                return
            
            # Update test type
            test.name = name
            test.category = category
            test.price = price
            test.description = description
            
            if self.db.update_test_type(test):
                messagebox.showinfo(_("Success"), _("Test updated successfully"))
                dialog.destroy()
                self.load_tests_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to update test"))
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text=_("Save"), command=save_test).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), 
                  command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Focus on first entry
        name_entry.focus()

    def delete_test(self):
        selected = self.tests_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a test"))
            return
        
        # Get the selected item
        item = self.tests_tree.selection()[0]
        
        # Try to get the test ID from tags first, fallback to values if needed
        try:
            test_id = self.tests_tree.item(item, "tags")[0]
        except (IndexError, TypeError):
            # If tags don't contain the ID, try to get it from the values
            values = self.tests_tree.item(item, "values")
            if values and len(values) > 0:
                test_id = values[0]  # ID is the first column
            else:
                messagebox.showerror(_("Error"), _("Unable to identify test"))
                return
        
        # Confirm deletion
        if messagebox.askyesno(_("Confirm Delete"), 
                              _("Are you sure you want to delete this test?")):
            # Try to delete by the ID we found
            if self.db.delete_test_type(test_id):
                messagebox.showinfo(_("Success"), _("Test deleted successfully"))
                self.load_tests_data()
            else:
                # If that fails, try to find the full ID first
                all_tests = self.db.get_all_test_types()
                full_test = next((t for t in all_tests if t.id.startswith(test_id) or t.id[:8] == test_id), None)
                
                if full_test and self.db.delete_test_type(full_test.id):
                    messagebox.showinfo(_("Success"), _("Test deleted successfully"))
                    self.load_tests_data()
                else:
                    messagebox.showerror(_("Error"), _("Failed to delete test"))
    
    def show_server_connection(self):
        """Show server connection dialog for admin users"""
        if not self.current_user or self.current_user.role != UserRole.ADMIN:
            messagebox.showerror(_("Error"), _("Access denied. Admin privileges required."))
            return
        
        # Create server connection dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Server Connection"))
        dialog.geometry("400x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Server configuration fields
        ttk.Label(dialog, text=_("Server Configuration"), font=("Arial", 14, "bold")).pack(pady=10)
        
        # Server URL
        ttk.Label(dialog, text=_("Server URL:"), foreground="#000080").pack(anchor=tk.W, padx=20, pady=(5, 0))
        server_url_entry = ttk.Entry(dialog, width=40)
        server_url_entry.pack(padx=20, pady=5)
        server_url_entry.insert(0, "https://example.com/api")  # Default value
        
        # API Key
        ttk.Label(dialog, text=_("API Key:"), foreground="#000080").pack(anchor=tk.W, padx=20, pady=(5, 0))
        api_key_entry = ttk.Entry(dialog, width=40, show="*")
        api_key_entry.pack(padx=20, pady=5)
        api_key_entry.insert(0, "your-api-key-here")  # Default value
        
        # Username
        ttk.Label(dialog, text=_("Username:"), foreground="#000080").pack(anchor=tk.W, padx=20, pady=(5, 0))
        username_entry = ttk.Entry(dialog, width=40)
        username_entry.pack(padx=20, pady=5)
        username_entry.insert(0, "admin")  # Default value
        
        # Password
        ttk.Label(dialog, text=_("Password:"), foreground="#000080").pack(anchor=tk.W, padx=20, pady=(5, 0))
        password_entry = ttk.Entry(dialog, width=40, show="*")
        password_entry.pack(padx=20, pady=5)
        password_entry.insert(0, "password")  # Default value
        
        def save_settings():
            # In a real application, this would save the server settings
            messagebox.showinfo(_("Success"), _("Server settings saved successfully"))
            dialog.destroy()
        
        def test_connection():
            # In a real application, this would test the server connection
            messagebox.showinfo(_("Test Connection"), _("Connection test functionality would be implemented here"))
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text=_("Save"), command=save_settings, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Test Connection"), command=test_connection, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), command=dialog.destroy, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)

    def change_admin_password(self):
        """Show change admin password dialog"""
        if not self.current_user or self.current_user.role != UserRole.ADMIN:
            messagebox.showerror(_("Error"), _("Access denied. Admin privileges required."))
            return
        
        # Create change password dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Change Admin Password"))
        dialog.geometry("400x250")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Password fields
        ttk.Label(dialog, text=_("Change Admin Password"), font=("Arial", 14, "bold")).pack(pady=10)
        
        # Current password
        ttk.Label(dialog, text=_("Current Password:"), foreground="#000080").pack(anchor=tk.W, padx=20, pady=(5, 0))
        current_password_entry = ttk.Entry(dialog, width=40, show="*")
        current_password_entry.pack(padx=20, pady=5)
        
        # New password
        ttk.Label(dialog, text=_("New Password:"), foreground="#000080").pack(anchor=tk.W, padx=20, pady=(5, 0))
        new_password_entry = ttk.Entry(dialog, width=40, show="*")
        new_password_entry.pack(padx=20, pady=5)
        
        # Confirm new password
        ttk.Label(dialog, text=_("Confirm New Password:"), foreground="#000080").pack(anchor=tk.W, padx=20, pady=(5, 0))
        confirm_password_entry = ttk.Entry(dialog, width=40, show="*")
        confirm_password_entry.pack(padx=20, pady=5)
        
        def save_password():
            current_password = current_password_entry.get()
            new_password = new_password_entry.get()
            confirm_password = confirm_password_entry.get()
            
            # Validation
            if not current_password or not new_password or not confirm_password:
                messagebox.showerror(_("Error"), _("Please fill in all password fields"))
                return
            
            if new_password != confirm_password:
                messagebox.showerror(_("Error"), _("New passwords do not match"))
                return
            
            if len(new_password) < 6:
                messagebox.showerror(_("Error"), _("Password must be at least 6 characters long"))
                return
            
            # Verify current password
            hashed_current = self.hash_password(current_password)
            if hashed_current != self.current_user.password_hash:
                messagebox.showerror(_("Error"), _("Current password is incorrect"))
                return
            
            # Update password
            self.current_user.password_hash = self.hash_password(new_password)
            if self.db.update_user(self.current_user):
                messagebox.showinfo(_("Success"), _("Password changed successfully"))
                dialog.destroy()
            else:
                messagebox.showerror(_("Error"), _("Failed to change password"))
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text=_("Save"), command=save_password, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), command=dialog.destroy, 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        
        # Focus on first entry
        current_password_entry.focus()

    def add_test(self):
        # Create add test dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Add New Test"))
        dialog.geometry("400x350")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Form fields
        ttk.Label(dialog, text=_("Test Name:")).pack(pady=5)
        name_entry = ttk.Entry(dialog, width=40)
        name_entry.pack(pady=5)
        
        ttk.Label(dialog, text=_("Category:")).pack(pady=5)
        category_entry = ttk.Entry(dialog, width=40)
        category_entry.pack(pady=5)
        
        ttk.Label(dialog, text=_("Price:")).pack(pady=5)
        price_entry = ttk.Entry(dialog, width=40)
        price_entry.pack(pady=5)
        
        ttk.Label(dialog, text=_("Description:")).pack(pady=5)
        desc_text = tk.Text(dialog, width=40, height=5)
        desc_text.pack(pady=5)
        
        def save_test():
            name = name_entry.get().strip()
            category = category_entry.get().strip()
            price_str = price_entry.get().strip()
            description = desc_text.get("1.0", tk.END).strip()
            
            # Validation
            if not name:
                messagebox.showerror(_("Error"), _("Please enter test name"))
                return
            
            if not category:
                messagebox.showerror(_("Error"), _("Please enter category"))
                return
            
            try:
                price = float(price_str)
                if price < 0:
                    raise ValueError()
            except ValueError:
                messagebox.showerror(_("Error"), _("Please enter a valid price"))
                return
            
            # Generate sequential three-digit test ID
            test_id = self.db.get_next_test_id()
            
            # Create test type with sequential ID
            test_type = TestType(
                id=test_id,
                name=name,
                description=description,
                price=price,
                category=category
            )
            
            if self.db.create_test_type(test_type):
                messagebox.showinfo(_("Success"), _("Test added successfully"))
                dialog.destroy()
                self.load_tests_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to add test"))
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text=_("Save"), command=save_test).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), 
                  command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Focus on first entry
        name_entry.focus()

    def view_test_details(self):
        selected = self.tests_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a test"))
            return
        
        # Get the selected item
        item = self.tests_tree.selection()[0]
        
        # Try to get the test ID from tags first, fallback to values if needed
        try:
            test_id = self.tests_tree.item(item, "tags")[0]
        except (IndexError, TypeError):
            # If tags don't contain the ID, try to get it from the values
            values = self.tests_tree.item(item, "values")
            if values and len(values) > 0:
                test_id = values[0]  # ID is the first column
            else:
                messagebox.showerror(_("Error"), _("Unable to identify test"))
                return
        
        # Get test details
        test = self.db.get_test_type(test_id)
        if not test:
            # Try to find test by display ID
            all_tests = self.db.get_all_test_types()
            test = next((t for t in all_tests if t.id.startswith(test_id) or t.id[:8] == test_id), None)
            
            if not test:
                messagebox.showerror(_("Error"), _("Test not found"))
                return
        
        # Create test details dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Test Details: {}").format(test.name))
        dialog.geometry("400x350")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Test details
        ttk.Label(dialog, text=_("Test Name: {}").format(test.name)).pack(pady=5)
        ttk.Label(dialog, text=_("Category: {}").format(test.category)).pack(pady=5)
        ttk.Label(dialog, text=_("Price: ${:.2f}").format(test.price)).pack(pady=5)
        ttk.Label(dialog, text=_("Description: {}").format(test.description)).pack(pady=5)
        
        # Close button
        ttk.Button(dialog, text=_("Close"), command=dialog.destroy).pack(pady=10)

    def edit_test(self):
        selected = self.tests_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a test"))
            return
        
        # Get the selected item
        item = self.tests_tree.selection()[0]
        
        # Try to get the test ID from tags first, fallback to values if needed
        try:
            test_id = self.tests_tree.item(item, "tags")[0]
        except (IndexError, TypeError):
            # If tags don't contain the ID, try to get it from the values
            values = self.tests_tree.item(item, "values")
            if values and len(values) > 0:
                test_id = values[0]  # ID is the first column
            else:
                messagebox.showerror(_("Error"), _("Unable to identify test"))
                return
        
        # Get test details
        test = self.db.get_test_type(test_id)
        if not test:
            # Try to find test by display ID
            all_tests = self.db.get_all_test_types()
            test = next((t for t in all_tests if t.id.startswith(test_id) or t.id[:8] == test_id), None)
            
            if not test:
                messagebox.showerror(_("Error"), _("Test not found"))
                return
        
        # Create edit test dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Edit Test: {}").format(test.name))
        dialog.geometry("400x350")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Form fields
        ttk.Label(dialog, text=_("Test Name:")).pack(pady=5)
        name_entry = ttk.Entry(dialog, width=40)
        name_entry.pack(pady=5)
        name_entry.insert(0, test.name)
        
        ttk.Label(dialog, text=_("Category:")).pack(pady=5)
        category_var = tk.StringVar()
        category_combo = ttk.Combobox(dialog, textvariable=category_var,
                                      values=[_("Blood"), _("Urine"), _("Biochemistry"), _("Imaging"), _("Genetics"), _("Hematology"), _("Microbiology"), _("Others")], state="readonly", width=37)
        category_combo.pack(pady=5)
        category_combo.set(test.category)
        
        ttk.Label(dialog, text=_("Price:")).pack(pady=5)
        price_entry = ttk.Entry(dialog, width=40)
        price_entry.pack(pady=5)
        price_entry.insert(0, str(test.price))
        
        ttk.Label(dialog, text=_("Description:")).pack(pady=5)
        description_entry = ttk.Entry(dialog, width=40)
        description_entry.pack(pady=5)
        description_entry.insert(0, test.description)
        
        def save_test():
            name = name_entry.get().strip()
            category = category_var.get().strip()
            price = price_entry.get().strip()
            description = description_entry.get().strip()
            
            # Validation
            if not name:
                messagebox.showerror(_("Error"), _("Please enter a test name"))
                return
            
            if not category:
                messagebox.showerror(_("Error"), _("Please select a category"))
                return
            
            if not price:
                messagebox.showerror(_("Error"), _("Please enter a price"))
                return
            
            try:
                price = float(price)
            except ValueError:
                messagebox.showerror(_("Error"), _("Price must be a number"))
                return
            
            if price < 0:
                messagebox.showerror(_("Error"), _("Price cannot be negative"))
                return
            
            # Update test type
            test.name = name
            test.category = category
            test.price = price
            test.description = description
            
            if self.db.update_test_type(test):
                messagebox.showinfo(_("Success"), _("Test updated successfully"))
                dialog.destroy()
                self.load_tests_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to update test"))
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text=_("Save"), command=save_test).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), 
                  command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Focus on first entry
        name_entry.focus()

    def delete_test(self):
        selected = self.tests_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a test"))
            return
        
        # Get the selected item
        item = self.tests_tree.selection()[0]
        
        # Try to get the test ID from tags first, fallback to values if needed
        try:
            test_id = self.tests_tree.item(item, "tags")[0]
        except (IndexError, TypeError):
            # If tags don't contain the ID, try to get it from the values
            values = self.tests_tree.item(item, "values")
            if values and len(values) > 0:
                test_id = values[0]  # ID is the first column
            else:
                messagebox.showerror(_("Error"), _("Unable to identify test"))
                return
        
        # Confirm deletion
        if messagebox.askyesno(_("Confirm Delete"), 
                              _("Are you sure you want to delete this test?")):
            # Try to delete by the ID we found
            if self.db.delete_test_type(test_id):
                messagebox.showinfo(_("Success"), _("Test deleted successfully"))
                self.load_tests_data()
            else:
                # If that fails, try to find the full ID first
                all_tests = self.db.get_all_test_types()
                full_test = next((t for t in all_tests if t.id.startswith(test_id) or t.id[:8] == test_id), None)
                
                if full_test and self.db.delete_test_type(full_test.id):
                    messagebox.showinfo(_("Success"), _("Test deleted successfully"))
                    self.load_tests_data()
                else:
                    messagebox.showerror(_("Error"), _("Failed to delete test"))
    
    def show_samples(self):
        self.current_screen = self.show_samples
        self.clear_content()
        
        # Samples header with 3D styling
        header_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("Sample Management"), 
                 style="Title.TLabel").pack(side=tk.LEFT)
        
        ttk.Button(header_frame, text=_("Add Sample"), 
                  command=self.add_sample, style="Accent.TButton").pack(side=tk.RIGHT)
        
        # Samples table with enhanced styling
        table_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        table_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create treeview with custom styling
        columns = (_("ID"), _("Barcode"), _("Test Request"), _("Collected At"), _("Status"))
        self.samples_tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        
        for col in columns:
            self.samples_tree.heading(col, text=col)
            self.samples_tree.column(col, width=120)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, 
                                 command=self.samples_tree.yview)
        self.samples_tree.configure(yscroll=scrollbar.set)
        
        self.samples_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load samples data
        self.load_samples_data()
        
        # Action buttons with 3D styling
        action_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        action_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(action_frame, text=_("View Details"), 
                  command=self.view_sample_details, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Update Status"), 
                  command=self.update_sample_status, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Generate Barcode"), 
                  command=self.generate_sample_barcode, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
    
    def load_samples_data(self):
        # Clear existing data
        for item in self.samples_tree.get_children():
            self.samples_tree.delete(item)
        
        # Load samples from database
        samples = self.db.get_all_samples()
        
        for sample in samples:
            # Get test request to get patient and test type info
            test_request = self.db.get_test_request(sample.test_request_id)
            patient_name = _("Unknown Patient")
            test_name = _("Unknown Test")
            
            if test_request:
                # Get patient info
                patient = self.db.get_patient(test_request.patient_id)
                if patient:
                    patient_name = patient.name
                
                # Get test type info
                test_type = self.db.get_test_type(test_request.test_type_id)
                if test_type:
                    test_name = test_type.name
            
            self.samples_tree.insert("", tk.END, values=(
                sample.id[:8],
                sample.barcode,
                f"{patient_name} - {test_name}",
                sample.collected_at.strftime("%Y-%m-%d %H:%M"),
                _(sample.status.value)
            ))

    def add_sample(self):
        # Create add sample dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Add New Sample"))
        dialog.geometry("500x400")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Form fields
        ttk.Label(dialog, text=_("Patient Name:")).pack(pady=5)
        patient_name_entry = ttk.Entry(dialog, width=50)
        patient_name_entry.pack(pady=5)
        
        # Test selection
        ttk.Label(dialog, text=_("Select Tests:")).pack(pady=5)
        test_frame = ttk.Frame(dialog)
        test_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Create listbox for available tests
        test_listbox = tk.Listbox(test_frame, selectmode=tk.MULTIPLE, height=6)
        test_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Add scrollbar
        test_scrollbar = ttk.Scrollbar(test_frame, orient=tk.VERTICAL, command=test_listbox.yview)
        test_listbox.configure(yscrollcommand=test_scrollbar.set)
        test_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load test types
        test_types = self.db.get_all_test_types()
        for test_type in test_types:
            test_listbox.insert(tk.END, test_type.name)
        
        # Selected tests display
        ttk.Label(dialog, text=_("Selected Tests:")).pack(pady=(10, 5))
        selected_tests_frame = ttk.Frame(dialog)
        selected_tests_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        selected_tests_listbox = tk.Listbox(selected_tests_frame, height=4)
        selected_tests_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        selected_tests_scrollbar = ttk.Scrollbar(selected_tests_frame, orient=tk.VERTICAL, 
                                                command=selected_tests_listbox.yview)
        selected_tests_listbox.configure(yscrollcommand=selected_tests_scrollbar.set)
        selected_tests_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Add and remove buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=5)
        
        def add_selected_tests():
            selections = test_listbox.curselection()
            for index in selections:
                test_name = test_listbox.get(index)
                if test_name not in selected_tests_listbox.get(0, tk.END):
                    selected_tests_listbox.insert(tk.END, test_name)
        
        def remove_selected_tests():
            selections = selected_tests_listbox.curselection()
            for index in reversed(selections):  # Remove in reverse order
                selected_tests_listbox.delete(index)
        
        ttk.Button(button_frame, text=_("Add Selected"), command=add_selected_tests).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Remove Selected"), command=remove_selected_tests).pack(side=tk.LEFT, padx=5)
        
        # Status
        ttk.Label(dialog, text=_("Status:")).pack(pady=5)
        status_var = tk.StringVar(value=_("Collected"))
        status_combo = ttk.Combobox(dialog, textvariable=status_var,
                                   values=[_("Collected"), _("Processing"), _("Completed")],
                                   state="readonly", width=47)
        status_combo.pack(pady=5)
        
        def save_sample():
            patient_name = patient_name_entry.get().strip()
            selected_tests = selected_tests_listbox.get(0, tk.END)
            status_text = status_var.get()
            
            # Validation
            if not patient_name:
                messagebox.showerror(_("Error"), _("Please enter patient name"))
                return
            
            if not selected_tests:
                messagebox.showerror(_("Error"), _("Please select at least one test"))
                return
            
            if not status_text:
                messagebox.showerror(_("Error"), _("Please select status"))
                return
            
            # Map status text to enum
            status_map = {
                _("Collected"): SampleStatus.COLLECTED,
                _("Processing"): SampleStatus.PROCESSING,
                _("Completed"): SampleStatus.COMPLETED
            }
            status = status_map.get(status_text, SampleStatus.COLLECTED)
            
            # For each selected test, create a test request and sample
            for test_name in selected_tests:
                # Find the test type
                test_type = None
                for t in test_types:
                    if t.name == test_name:
                        test_type = t
                        break
                
                if test_type:
                    # Create a patient (in a real app, you'd look up existing patient)
                    patient = Patient(
                        id=self.db.generate_patient_id(),
                        name=patient_name,
                        age=0,  # Default age
                        gender=Gender.OTHER,  # Default gender
                        contact_info=""  # Default contact
                    )
                    self.db.create_patient(patient)
                    
                    # Create test request
                    test_request = TestRequest(
                        id=str(uuid.uuid4()),
                        patient_id=patient.id,
                        test_type_id=test_type.id,
                        requested_by=self.current_user.username if self.current_user else "System",
                        requested_at=datetime.now(),
                        status=TestStatus.PENDING
                    )
                    self.db.create_test_request(test_request)
                    
                    # Generate barcode
                    barcode = generate_barcode()
                    
                    # Create sample
                    sample = Sample(
                        id=str(uuid.uuid4()),
                        test_request_id=test_request.id,
                        barcode=barcode,
                        collected_at=datetime.now(),
                        status=status,
                        notes=f"Sample for {test_name}"
                    )
                    self.db.create_sample(sample)
            
            messagebox.showinfo(_("Success"), _("Samples added successfully"))
            dialog.destroy()
            self.load_samples_data()
        
        # Buttons
        save_button_frame = ttk.Frame(dialog)
        save_button_frame.pack(pady=20)
        
        ttk.Button(save_button_frame, text=_("Save"), command=save_sample).pack(side=tk.LEFT, padx=5)
        ttk.Button(save_button_frame, text=_("Cancel"), 
                  command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Focus on first entry
        patient_name_entry.focus()
    
    def view_sample_details(self):
        selected = self.samples_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a sample"))
            return
        
        sample_id = self.samples_tree.item(selected[0])['values'][0]
        sample = self.db.get_sample_by_id(sample_id)
        if not sample:
            messagebox.showerror(_("Error"), _("Sample not found"))
            return
        
        test_request = self.db.get_test_request(sample.test_request_id)
        if not test_request:
            messagebox.showerror(_("Error"), _("Test request not found"))
            return
        
        patient = self.db.get_patient(test_request.patient_id)
        if not patient:
            messagebox.showerror(_("Error"), _("Patient not found"))
            return
        
        test_type = self.db.get_test_type(test_request.test_type_id)
        if not test_type:
            messagebox.showerror(_("Error"), _("Test type not found"))
            return
        
        # Create detail dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Sample Details"))
        dialog.geometry("500x400")
        dialog.transient(self.root)
        dialog.grab_set()
        
        ttk.Label(dialog, text=_("Sample ID:")).pack(pady=5)
        ttk.Label(dialog, text=f"{sample.id[:8]}").pack(pady=5)
        
        ttk.Label(dialog, text=_("Barcode:")).pack(pady=5)
        ttk.Label(dialog, text=sample.barcode).pack(pady=5)
        
        ttk.Label(dialog, text=_("Patient Name:")).pack(pady=5)
        ttk.Label(dialog, text=patient.name).pack(pady=5)
        
        ttk.Label(dialog, text=_("Test Name:")).pack(pady=5)
        ttk.Label(dialog, text=test_type.name).pack(pady=5)
        
        ttk.Label(dialog, text=_("Collected At:")).pack(pady=5)
        ttk.Label(dialog, text=sample.collected_at.strftime("%Y-%m-%d %H:%M")).pack(pady=5)
        
        ttk.Label(dialog, text=_("Status:")).pack(pady=5)
        ttk.Label(dialog, text=_(sample.status.value)).pack(pady=5)
    
    def update_sample_status(self):
        selected = self.samples_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a sample"))
            return
        
        sample_id = self.samples_tree.item(selected[0])['values'][0]
        sample = self.db.get_sample_by_id(sample_id)
        if not sample:
            messagebox.showerror(_("Error"), _("Sample not found"))
            return
        
        status_var = tk.StringVar(value=_(sample.status.value))
        
        # Create update dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Update Sample Status"))
        dialog.geometry("300x150")
        dialog.transient(self.root)
        dialog.grab_set()
        
        ttk.Label(dialog, text=_("Status:")).pack(pady=5)
        status_combo = ttk.Combobox(dialog, textvariable=status_var,
                                   values=[_("Collected"), _("Processing"), _("Completed")],
                                   state="readonly", width=37)
        status_combo.pack(pady=5)
        
        def update_status():
            status_text = status_var.get()
            
            # Map status text to enum
            status_map = {
                _("Collected"): SampleStatus.COLLECTED,
                _("Processing"): SampleStatus.PROCESSING,
                _("Completed"): SampleStatus.COMPLETED
            }
            status = status_map.get(status_text, SampleStatus.COLLECTED)
            
            if self.db.update_sample_status(sample.id, status):
                messagebox.showinfo(_("Success"), _("Sample status updated successfully"))
                dialog.destroy()
                self.load_samples_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to update sample status"))
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text=_("Update"), command=update_status).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), 
                  command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def generate_sample_barcode(self):
        selected = self.samples_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a sample"))
            return
        
        sample_id = self.samples_tree.item(selected[0])['values'][0]
        sample = self.db.get_sample_by_id(sample_id)
        if not sample:
            messagebox.showerror(_("Error"), _("Sample not found"))
            return
        
        # Generate barcode
        barcode = generate_barcode()
        
        if self.db.update_sample_barcode(sample.id, barcode):
            messagebox.showinfo(_("Success"), _("Barcode generated successfully"))
            self.load_samples_data()
        else:
            messagebox.showerror(_("Error"), _("Failed to generate barcode"))
    
    def show_results(self):
        self.current_screen = self.show_results
        self.clear_content()
        
        # Professional Results Management header with enhanced 3D styling
        header_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        # Title with professional styling
        title_frame = ttk.Frame(header_frame, style="Card.TFrame")
        title_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(title_frame, text=_("Professional Results Management"), 
                 style="Title.TLabel", font=("Arial", 16, "bold")).pack(side=tk.LEFT)
        
        # Action buttons with improved styling
        button_frame = ttk.Frame(header_frame, style="Card.TFrame")
        button_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(button_frame, text=_("➕ Create New Result"), 
                  command=self.create_new_result, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("📋 Manage Templates"), 
                  command=self.manage_test_templates, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("🖨️ Print Selected"), 
                  command=self.print_selected_result, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        
        # Enhanced results table with professional styling
        table_container = ttk.LabelFrame(self.content_frame, text=_("Medical Test Results"), 
                                        style="Card.TFrame", padding=10)
        table_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create treeview with professional styling
        columns = (_("Result ID"), _("Patient Name"), _("Test Type"), _("Status"), _("Created Date"), _("Signed By"))
        self.results_tree = ttk.Treeview(table_container, columns=columns, show="headings", height=15)
        
        # Configure column headings with better styling
        for col in columns:
            self.results_tree.heading(col, text=col, anchor=tk.CENTER)
            self.results_tree.column(col, width=150, anchor=tk.CENTER)
        
        # Configure specific column widths
        self.results_tree.column(columns[0], width=100)  # Result ID
        self.results_tree.column(columns[1], width=180)  # Patient Name
        self.results_tree.column(columns[2], width=150)  # Test Type
        self.results_tree.column(columns[3], width=100)  # Status
        self.results_tree.column(columns[4], width=150)  # Created Date
        self.results_tree.column(columns[5], width=150)  # Signed By
        
        # Add scrollbar with professional styling
        scrollbar_v = ttk.Scrollbar(table_container, orient=tk.VERTICAL, command=self.results_tree.yview)
        scrollbar_h = ttk.Scrollbar(table_container, orient=tk.HORIZONTAL, command=self.results_tree.xview)
        self.results_tree.configure(yscrollcommand=scrollbar_v.set, xscrollcommand=scrollbar_h.set)
        
        # Pack treeview and scrollbars
        self.results_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar_v.pack(side=tk.RIGHT, fill=tk.Y)
        scrollbar_h.pack(side=tk.BOTTOM, fill=tk.X)
        
        # Load results data
        self.load_results_data()
        
        # Enhanced action buttons with icons
        action_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        action_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(action_frame, text=_("👁️ View Details"), 
                  command=self.view_result_details, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("✏️ Edit Result"), 
                  command=self.edit_result, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("🗑️ Delete Result"), 
                  command=self.delete_result, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("🖨️ Print Result"), 
                  command=self.print_selected_result, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        
        # Bind double-click to view details
        self.results_tree.bind("<Double-1>", lambda event: self.view_result_details())
    
    def load_results_data(self):
        # Check if results_tree exists
        if not hasattr(self, 'results_tree'):
            return
            
        # Clear existing data
        for item in self.results_tree.get_children():
            self.results_tree.delete(item)
        
        # Load medical reports from database
        results = self.db.get_all_medical_reports()
        
        for result in results:
            test_request = self.db.get_test_request(result.test_request_id)
            patient = None
            test_type = None
            status = _("Pending")
            
            if test_request:
                patient = self.db.get_patient(test_request.patient_id)
                test_type = self.db.get_test_type(test_request.test_type_id)
                status = _(test_request.status.value)
            
            patient_name = _("Unknown Patient")
            test_name = _("Unknown Test")
            
            if patient:
                patient_name = patient.name
            
            if test_type:
                test_name = test_type.name
            
            # Check the number of columns in the current results_tree
            columns = self.results_tree["columns"]
            if len(columns) == 6:
                # show_results function with 6 columns
                self.results_tree.insert("", tk.END, values=(
                    result.id[:8],
                    patient_name,
                    test_name,
                    status,
                    result.created_at.strftime("%Y-%m-%d %H:%M") if result.created_at else _("Unknown"),
                    result.signed_by
                ))
            elif len(columns) == 5:
                # dashboard with 5 columns
                self.results_tree.insert("", tk.END, values=(
                    result.id[:8],
                    patient_name,
                    test_name,
                    status,
                    result.created_at.strftime("%Y-%m-%d %H:%M") if result.created_at else _("Unknown")
                ))
    
    def create_new_result(self):
        """Create a new medical result with patient selection and multiple test selection"""
        # Create result dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Create New Medical Result"))
        dialog.geometry("600x500")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Patient selection
        patient_frame = ttk.LabelFrame(dialog, text=_("Select Patient"), padding=10)
        patient_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Patient combobox
        ttk.Label(patient_frame, text=_("Patient:")).pack(anchor=tk.W)
        patient_var = tk.StringVar()
        patient_combo = ttk.Combobox(patient_frame, textvariable=patient_var, state="readonly", width=50)
        patient_combo.pack(fill=tk.X, pady=5)
        
        # Load patients
        patients = self.db.get_all_patients()
        patient_names = [f"{p.name} (ID: {p.id})" for p in patients]
        patient_map = {f"{p.name} (ID: {p.id})": p.id for p in patients}
        patient_combo['values'] = patient_names
        
        # Tests selection
        tests_frame = ttk.LabelFrame(dialog, text=_("Select Tests Performed"), padding=10)
        tests_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Available tests listbox
        ttk.Label(tests_frame, text=_("Available Tests:")).pack(anchor=tk.W)
        available_frame = ttk.Frame(tests_frame)
        available_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        available_listbox = tk.Listbox(available_frame, selectmode=tk.MULTIPLE, height=6)
        available_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        available_scrollbar = ttk.Scrollbar(available_frame, orient=tk.VERTICAL, command=available_listbox.yview)
        available_listbox.configure(yscrollcommand=available_scrollbar.set)
        available_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Selected tests listbox
        ttk.Label(tests_frame, text=_("Selected Tests:")).pack(anchor=tk.W, pady=(10, 0))
        selected_frame = ttk.Frame(tests_frame)
        selected_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        selected_listbox = tk.Listbox(selected_frame, height=4)
        selected_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        selected_scrollbar = ttk.Scrollbar(selected_frame, orient=tk.VERTICAL, command=selected_listbox.yview)
        selected_listbox.configure(yscrollcommand=selected_scrollbar.set)
        selected_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load test types
        test_types = self.db.get_all_test_types()
        test_map = {t.name: t for t in test_types}
        for test_type in test_types:
            available_listbox.insert(tk.END, test_type.name)
        
        # Add and remove buttons
        button_frame = ttk.Frame(tests_frame)
        button_frame.pack(pady=5)
        
        def add_selected_tests():
            selections = available_listbox.curselection()
            for index in selections:
                test_name = available_listbox.get(index)
                if test_name not in selected_listbox.get(0, tk.END):
                    selected_listbox.insert(tk.END, test_name)
        
        def remove_selected_tests():
            selections = selected_listbox.curselection()
            for index in reversed(selections):  # Remove in reverse order
                selected_listbox.delete(index)
        
        ttk.Button(button_frame, text=_("Add Selected"), command=add_selected_tests).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Remove Selected"), command=remove_selected_tests).pack(side=tk.LEFT, padx=5)
        
        # Template loading section
        template_frame = ttk.LabelFrame(dialog, text=_("Load Template"), padding=10)
        template_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Template selection
        template_var = tk.StringVar()
        ttk.Label(template_frame, text=_("Select Template:")).pack(anchor=tk.W)
        
        # Get all templates
        templates = self.db.get_all_test_templates()
        template_ids = [template.id for template in templates]
        
        template_combo = ttk.Combobox(template_frame, textvariable=template_var, 
                                     values=template_ids, state="readonly", width=50)
        template_combo.pack(fill=tk.X, pady=5)
        
        # Load from Word file button
        def load_from_word():
            try:
                file_path = filedialog.askopenfilename(
                    title=_("Select Word Template File"),
                    filetypes=[(_("Word files"), "*.docx"), (_("All files"), "*.*")]
                )
                
                if file_path:
                    # Load content from Word document
                    doc = Document(file_path)
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    content_text.delete("1.0", tk.END)
                    content_text.insert("1.0", content)
                    messagebox.showinfo(_("Success"), _("Template loaded successfully"))
            except Exception as e:
                messagebox.showerror(_("Error"), f"{_('Failed to load template')}: {str(e)}")
        
        # Template buttons
        template_button_frame = ttk.Frame(template_frame)
        template_button_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(template_button_frame, text=_("Load from Word File"), 
                  command=load_from_word, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        
        # Apply template button
        def apply_template():
            selected_template_id = template_var.get()
            if not selected_template_id:
                messagebox.showwarning(_("Warning"), _("Please select a template"))
                return
            
            # Get template content
            template = self.db.get_test_template(selected_template_id)
            if template:
                content_text.delete("1.0", tk.END)
                content_text.insert("1.0", template.template_content)
            else:
                messagebox.showerror(_("Error"), _("Template not found"))
        
        ttk.Button(template_button_frame, text=_("Apply Template"), 
                  command=apply_template, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        
        # Result content
        content_frame = ttk.LabelFrame(dialog, text=_("Result Content"), padding=10)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        content_text = tk.Text(content_frame, wrap=tk.WORD, height=10)
        content_text.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Scrollbar for content text
        content_scrollbar = ttk.Scrollbar(content_text, orient=tk.VERTICAL, command=content_text.yview)
        content_text.configure(yscrollcommand=content_scrollbar.set)
        content_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        def save_result():
            patient_name = patient_var.get()
            selected_tests = selected_listbox.get(0, tk.END)
            content = content_text.get("1.0", tk.END).strip()
            
            # Validation
            if not patient_name:
                messagebox.showerror(_("Error"), _("Please select a patient"))
                return
            
            if not selected_tests:
                messagebox.showerror(_("Error"), _("Please select at least one test"))
                return
            
            if not content:
                messagebox.showerror(_("Error"), _("Please enter result content"))
                return
            
            # Get patient ID
            patient_id = patient_map.get(patient_name)
            if not patient_id:
                messagebox.showerror(_("Error"), _("Invalid patient selection"))
                return
            
            # For each selected test, create a test request and medical report
            for test_name in selected_tests:
                test_type = test_map.get(test_name)
                if test_type:
                    # Create test request
                    test_request = TestRequest(
                        id=str(uuid.uuid4()),
                        patient_id=patient_id,
                        test_type_id=test_type.id,
                        requested_by=self.current_user.username if self.current_user else "System",
                        requested_at=datetime.now(),
                        status=TestStatus.COMPLETED
                    )
                    self.db.create_test_request(test_request)
                    
                    # Create medical report
                    report = MedicalReport(
                        id=str(uuid.uuid4()),
                        test_request_id=test_request.id,
                        content=content,  # Same content for all tests in this implementation
                        signed_by=self.current_user.id if self.current_user else "N/A",
                        signed_at=datetime.now()
                    )
                    self.db.create_medical_report(report)
            
            messagebox.showinfo(_("Success"), _("Results saved successfully"))
            dialog.destroy()
            self.load_results_data()
        
        ttk.Button(button_frame, text=_("Save Result"), command=save_result, style="Accent.TButton").pack(side=tk.RIGHT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), command=dialog.destroy, style="Accent.TButton").pack(side=tk.RIGHT, padx=5)
        
        # Focus on first entry
        patient_combo.focus()

    def create_new_result(self):
        """Create a new medical result with patient selection and test-specific templates"""
        # Create professional result dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("🏥 Create Professional Medical Result"))
        dialog.geometry("800x700")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Create notebook for better organization
        notebook = ttk.Notebook(dialog)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Patient and test selection tab
        selection_frame = ttk.Frame(notebook)
        notebook.add(selection_frame, text=_("📋 Patient & Test Selection"))
        
        # Patient selection
        patient_frame = ttk.LabelFrame(selection_frame, text=_("👤 Select Patient"), padding=15)
        patient_frame.pack(fill=tk.X, padx=15, pady=10)
        
        # Patient combobox
        ttk.Label(patient_frame, text=_("Patient:"), font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
        patient_var = tk.StringVar()
        patient_combo = ttk.Combobox(patient_frame, textvariable=patient_var, state="readonly", width=60, font=("Arial", 10))
        patient_combo.pack(fill=tk.X, pady=5)
        
        # Load patients
        patients = self.db.get_all_patients()
        patient_names = [f"{p.name} (ID: {p.id})" for p in patients]
        patient_map = {f"{p.name} (ID: {p.id})": p.id for p in patients}
        patient_combo['values'] = patient_names
        
        # Test selection with template preview
        tests_frame = ttk.LabelFrame(selection_frame, text=_("🧪 Select Test and Apply Template"), padding=15)
        tests_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)
        
        # Test selection
        ttk.Label(tests_frame, text=_("Select Test:"), font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
        test_var = tk.StringVar()
        test_combo = ttk.Combobox(tests_frame, textvariable=test_var, state="readonly", width=60, font=("Arial", 10))
        test_combo.pack(fill=tk.X, pady=5)
        
        # Load test types
        test_types = self.db.get_all_test_types()
        test_map = {t.name: t for t in test_types}
        test_combo['values'] = [t.name for t in test_types]
        
        # Template preview section
        template_preview_frame = ttk.LabelFrame(tests_frame, text=_("📄 Template Preview"), padding=10)
        template_preview_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=10)
        
        template_preview_text = tk.Text(template_preview_frame, wrap=tk.WORD, height=10, font=("Arial", 10), state=tk.DISABLED)
        template_preview_text.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Scrollbar for template preview
        template_preview_scrollbar = ttk.Scrollbar(template_preview_text, orient=tk.VERTICAL, 
                                              command=template_preview_text.yview)
        template_preview_text.configure(yscrollcommand=template_preview_scrollbar.set)
        template_preview_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load template when test is selected
        def on_test_selected(event=None):
            test_name = test_var.get()
            if test_name:
                test_type = test_map.get(test_name)
                if test_type:
                    # Get template for this test type
                    template = self.db.get_test_template_by_test_type(test_type.id)
                    if template:
                        template_preview_text.config(state=tk.NORMAL)
                        template_preview_text.delete("1.0", tk.END)
                        template_preview_text.insert("1.0", template.template_content)
                        template_preview_text.config(state=tk.DISABLED)
                    else:
                        template_preview_text.config(state=tk.NORMAL)
                        template_preview_text.delete("1.0", tk.END)
                        template_preview_text.insert("1.0", _("No template found for this test type. You can create one in the template management section."))
                        template_preview_text.config(state=tk.DISABLED)
    
        test_combo.bind("<<ComboboxSelected>>", on_test_selected)
        
        # Result content tab
        content_frame = ttk.Frame(notebook)
        notebook.add(content_frame, text=_("📝 Result Content"))
        
        # Result content with template application
        content_main_frame = ttk.LabelFrame(content_frame, text=_("📋 Result Content"), padding=15)
        content_main_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)
        
        # Template management buttons
        template_button_frame = ttk.Frame(content_main_frame)
        template_button_frame.pack(fill=tk.X, pady=(0, 10))
        
        def load_from_word():
            try:
                file_path = filedialog.askopenfilename(
                    title=_("Select Word Template File"),
                    filetypes=[(_("Word files"), "*.docx"), (_("All files"), "*.*")]
                )
                
                if file_path:
                    # Load content from Word document
                    doc = Document(file_path)
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    content_text.delete("1.0", tk.END)
                    content_text.insert("1.0", content)
                    messagebox.showinfo(_("Success"), _("Template loaded successfully from Word file"))
            except Exception as e:
                messagebox.showerror(_("Error"), f"{_('Failed to load template')}: {str(e)}")
    
        def apply_selected_template():
            test_name = test_var.get()
            if not test_name:
                messagebox.showwarning(_("Warning"), _("Please select a test first"))
                return
            
            test_type = test_map.get(test_name)
            if not test_type:
                messagebox.showerror(_("Error"), _("Invalid test selection"))
                return
            
            # Get template for this test type
            template = self.db.get_test_template_by_test_type(test_type.id)
            if template:
                content_text.delete("1.0", tk.END)
                content_text.insert("1.0", template.template_content)
                messagebox.showinfo(_("Success"), _("Template applied successfully"))
            else:
                messagebox.showinfo(_("Info"), _("No template found for this test type. Create one in template management."))
    
        ttk.Button(template_button_frame, text=_("📂 Load from Word File"), 
              command=load_from_word, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(template_button_frame, text=_("📋 Apply Test Template"), 
              command=apply_selected_template, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
    
        # Content text area with enhanced styling
        ttk.Label(content_main_frame, text=_("Result Content:"), font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
        content_text = tk.Text(content_main_frame, wrap=tk.WORD, height=20, font=("Arial", 11))
        content_text.pack(fill=tk.BOTH, expand=True, pady=5)
    
        # Scrollbar for content text
        content_scrollbar = ttk.Scrollbar(content_main_frame, orient=tk.VERTICAL, command=content_text.yview)
        content_text.configure(yscrollcommand=content_scrollbar.set)
        content_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    
        # Buttons with professional styling
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
    
        def save_result():
            patient_name = patient_var.get()
            test_name = test_var.get()
            content = content_text.get("1.0", tk.END).strip()
            
            # Validation
            if not patient_name:
                messagebox.showerror(_("Error"), _("Please select a patient"))
                return
            
            if not test_name:
                messagebox.showerror(_("Error"), _("Please select a test"))
                return
            
            if not content:
                messagebox.showerror(_("Error"), _("Please enter result content"))
                return
            
            # Get patient ID
            patient_id = patient_map.get(patient_name)
            if not patient_id:
                messagebox.showerror(_("Error"), _("Invalid patient selection"))
                return
            
            # Get test type
            test_type = test_map.get(test_name)
            if not test_type:
                messagebox.showerror(_("Error"), _("Invalid test selection"))
                return
            
            # Create test request
            test_request = TestRequest(
                id=str(uuid.uuid4()),
                patient_id=patient_id,
                test_type_id=test_type.id,
                requested_by=self.current_user.username if self.current_user else "System",
                requested_at=datetime.now(),
                status=TestStatus.COMPLETED
            )
            self.db.create_test_request(test_request)
            
            # Create medical report
            report = MedicalReport(
                id=str(uuid.uuid4()),
                test_request_id=test_request.id,
                content=content,
                signed_by=self.current_user.id if self.current_user else "N/A",
                signed_at=datetime.now()
            )
            self.db.create_medical_report(report)
            
            messagebox.showinfo(_("Success"), _("Medical result saved successfully"))
            dialog.destroy()
            self.load_results_data()
    
        ttk.Button(button_frame, text=_("💾 Save Result"), command=save_result, style="Accent.TButton").pack(side=tk.RIGHT, padx=5)
        ttk.Button(button_frame, text=_("❌ Cancel"), command=dialog.destroy, style="Accent.TButton").pack(side=tk.RIGHT, padx=5)
    
        # Focus on first entry
        patient_combo.focus()

    def view_result_details(self):
        """View details of a selected medical result"""
        selected = self.results_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a result"))
            return
        
        # Get the selected result ID from the item's tags
        item = self.results_tree.selection()[0]
        report_id = self.results_tree.item(item, "tags")[0]
        
        # Get report details from database
        selected_report = self.db.get_medical_report(report_id)
        
        if not selected_report:
            messagebox.showerror(_("Error"), _("Result not found"))
            return
        
        # Get associated test request, patient, and test type
        test_request = self.db.get_test_request(selected_report.test_request_id)
        if not test_request:
            messagebox.showerror(_("Error"), _("Test request not found"))
            return
            
        patient = self.db.get_patient(test_request.patient_id)
        if not patient:
            messagebox.showerror(_("Error"), _("Patient not found"))
            return
            
        test_type = self.db.get_test_type(test_request.test_type_id)
        if not test_type:
            messagebox.showerror(_("Error"), _("Test type not found"))
            return
        
        # Create detail dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Result Details"))
        dialog.geometry("600x550")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Patient info
        patient_frame = ttk.LabelFrame(dialog, text=_("Patient Information"), padding=10)
        patient_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(patient_frame, text=f"{_('Name')}: {patient.name}").pack(anchor=tk.W)
        ttk.Label(patient_frame, text=f"{_('ID')}: {patient.id}").pack(anchor=tk.W)
        ttk.Label(patient_frame, text=f"{_('Age')}: {patient.age}").pack(anchor=tk.W)
        ttk.Label(patient_frame, text=f"{_('Gender')}: {_(patient.gender.value)}").pack(anchor=tk.W)
        
        # Test info
        test_frame = ttk.LabelFrame(dialog, text=_("Test Information"), padding=10)
        test_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(test_frame, text=f"{_('Test Type')}: {test_type.name}").pack(anchor=tk.W)
        ttk.Label(test_frame, text=f"{_('Test ID')}: {test_request.id}").pack(anchor=tk.W)
        ttk.Label(test_frame, text=f"{_('Status')}: {_(test_request.status.value)}").pack(anchor=tk.W)
        
        # Result content
        result_frame = ttk.LabelFrame(dialog, text=_("Result Content"), padding=10)
        result_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        result_text = tk.Text(result_frame, wrap=tk.WORD)
        result_text.insert("1.0", selected_report.content)
        result_text.config(state=tk.DISABLED)  # Make it read-only
        result_text.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Scrollbar for result text
        scrollbar = ttk.Scrollbar(result_text, orient=tk.VERTICAL, command=result_text.yview)
        result_text.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(button_frame, text=_("Print"), 
                  command=lambda: self.print_result_from_details(selected_report, patient, test_type, test_request), 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Close"), 
                  command=dialog.destroy, 
                  style="Accent.TButton").pack(side=tk.RIGHT, padx=5)

    def print_result_from_details(self, report, patient, test_type, test_request):
        """Print result from the details view"""
        self.create_result_print_dialog(report, patient, test_type, test_request)

    def edit_result(self):
        """Edit an existing medical result with template support"""
        selected = self.results_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a result to edit"))
            return
        
        # Get the selected result ID from the item's tags
        item = self.results_tree.selection()[0]
        report_id = self.results_tree.item(item, "tags")[0]
        
        # Get report details from database
        selected_report = self.db.get_medical_report(report_id)
        
        if not selected_report:
            messagebox.showerror(_("Error"), _("Result not found"))
            return
        
        # Get associated test request, patient, and test type
        test_request = self.db.get_test_request(selected_report.test_request_id)
        if not test_request:
            messagebox.showerror(_("Error"), _("Test request not found"))
            return
            
        patient = self.db.get_patient(test_request.patient_id)
        if not patient:
            messagebox.showerror(_("Error"), _("Patient not found"))
            return
            
        test_type = self.db.get_test_type(test_request.test_type_id)
        if not test_type:
            messagebox.showerror(_("Error"), _("Test type not found"))
            return
    
        # Create professional edit dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("✏️ Edit Medical Result"))
        dialog.geometry("800x700")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Professional header
        header_frame = ttk.Frame(dialog, style="Header.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("✏️ Edit Medical Test Result"), 
                 font=("Arial", 16, "bold"), style="Header.TLabel").pack()
        
        # Create notebook for better organization
        notebook = ttk.Notebook(dialog)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Patient and test info tab
        info_frame = ttk.Frame(notebook)
        notebook.add(info_frame, text=_("📋 Patient & Test Info"))
        
        # Patient info with enhanced styling
        patient_frame = ttk.LabelFrame(info_frame, text=_("👤 Patient Information"), padding=15)
        patient_frame.pack(fill=tk.X, padx=15, pady=10)
        
        ttk.Label(patient_frame, text=f"{_('Patient')}: {patient.name}", font=("Arial", 11)).pack(anchor=tk.W, pady=2)
        ttk.Label(patient_frame, text=f"{_('Patient ID')}: {patient.id}", font=("Arial", 11)).pack(anchor=tk.W, pady=2)
        ttk.Label(patient_frame, text=f"{_('Age')}: {patient.age}", font=("Arial", 11)).pack(anchor=tk.W, pady=2)
        ttk.Label(patient_frame, text=f"{_('Gender')}: {_(patient.gender.value)}", font=("Arial", 11)).pack(anchor=tk.W, pady=2)
        
        # Test info with enhanced styling
        test_frame = ttk.LabelFrame(info_frame, text=_("🧪 Test Information"), padding=15)
        test_frame.pack(fill=tk.X, padx=15, pady=10)
        
        ttk.Label(test_frame, text=f"{_('Test Type')}: {test_type.name}", font=("Arial", 11)).pack(anchor=tk.W, pady=2)
        ttk.Label(test_frame, text=f"{_('Test Category')}: {test_type.category}", font=("Arial", 11)).pack(anchor=tk.W, pady=2)
        ttk.Label(test_frame, text=f"{_('Test ID')}: {test_request.id}", font=("Arial", 11)).pack(anchor=tk.W, pady=2)
        ttk.Label(test_frame, text=f"{_('Status')}: {_(test_request.status.value)}", font=("Arial", 11)).pack(anchor=tk.W, pady=2)
        ttk.Label(test_frame, text=f"{_('Requested By')}: {test_request.requested_by}", font=("Arial", 11)).pack(anchor=tk.W, pady=2)
        ttk.Label(test_frame, text=f"{_('Requested At')}: {test_request.requested_at.strftime('%Y-%m-%d %H:%M')}", font=("Arial", 11)).pack(anchor=tk.W, pady=2)
        
        # Template management tab
        template_frame = ttk.Frame(notebook)
        notebook.add(template_frame, text=_("📄 Template Management"))
        
        # Template loading section
        template_load_frame = ttk.LabelFrame(template_frame, text=_("📥 Load Template"), padding=15)
        template_load_frame.pack(fill=tk.X, padx=15, pady=10)
        
        # Template selection
        ttk.Label(template_load_frame, text=_("Select Template:"), font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
        
        # Get templates for this test type
        template = self.db.get_test_template_by_test_type(test_type.id)
        templates = []
        if template:
            templates.append(template.id)
    
        template_var = tk.StringVar()
        template_combo = ttk.Combobox(template_load_frame, textvariable=template_var, 
                                 values=templates, state="readonly", width=60, font=("Arial", 10))
        template_combo.pack(fill=tk.X, pady=5)
        
        # Template buttons with enhanced styling
        template_button_frame = ttk.Frame(template_load_frame)
        template_button_frame.pack(fill=tk.X, pady=10)
        
        def load_from_word():
            try:
                file_path = filedialog.askopenfilename(
                    title=_("Select Word Template File"),
                    filetypes=[(_("Word files"), "*.docx"), (_("All files"), "*.*")]
                )
                
                if file_path:
                    # Load content from Word document
                    doc = Document(file_path)
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    content_text.delete("1.0", tk.END)
                    content_text.insert("1.0", content)
                    messagebox.showinfo(_("Success"), _("Template loaded successfully from Word file"))
            except Exception as e:
                messagebox.showerror(_("Error"), f"{_('Failed to load template')}: {str(e)}")
    
        def apply_template():
            selected_template_id = template_var.get()
            if not selected_template_id:
                messagebox.showwarning(_("Warning"), _("Please select a template"))
                return
            
            # Get template content
            template = self.db.get_test_template(selected_template_id)
            if template:
                content_text.delete("1.0", tk.END)
                content_text.insert("1.0", template.template_content)
                messagebox.showinfo(_("Success"), _("Template applied successfully"))
            else:
                messagebox.showerror(_("Error"), _("Template not found"))
    
        ttk.Button(template_button_frame, text=_("📂 Load from Word File"), 
              command=load_from_word, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(template_button_frame, text=_("📋 Apply Template"), 
              command=apply_template, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        
        # Save template section
        template_save_frame = ttk.LabelFrame(template_frame, text=_("💾 Save as Template"), padding=15)
        template_save_frame.pack(fill=tk.X, padx=15, pady=10)
        
        def save_template():
            try:
                # Get current content
                content = content_text.get("1.0", tk.END).strip()
                
                if not content:
                    messagebox.showwarning(_("Warning"), _("Please enter template content"))
                    return
                
                # Check if template already exists for this test type
                existing_template = self.db.get_test_template_by_test_type(test_type.id)
                
                if existing_template:
                    # Update existing template
                    existing_template.template_content = content
                    existing_template.updated_at = datetime.now()
                    if self.db.update_test_template(existing_template):
                        messagebox.showinfo(_("Success"), _("Template updated successfully"))
                    else:
                        messagebox.showerror(_("Error"), _("Failed to update template"))
                else:
                    # Create new template
                    new_template = TestTemplate(
                        id=str(uuid.uuid4()),
                        test_type_id=test_type.id,
                        template_content=content
                    )
                    if self.db.create_test_template(new_template):
                        messagebox.showinfo(_("Success"), _("Template saved successfully"))
                        # Update combo box
                        template_combo['values'] = list(template_combo['values']) + [new_template.id]
                    else:
                        messagebox.showerror(_("Error"), _("Failed to save template"))
            except Exception as e:
                messagebox.showerror(_("Error"), f"{_('Failed to save template')}: {str(e)}")
    
        ttk.Button(template_save_frame, text=_("💾 Save Current Content as Template"), 
              command=save_template, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        
        # Result content tab
        content_frame = ttk.Frame(notebook)
        notebook.add(content_frame, text=_("📝 Result Content"))
        
        # Result content with enhanced styling
        content_main_frame = ttk.LabelFrame(content_frame, text=_("📋 Edit Result Content"), padding=15)
        content_main_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)
        
        ttk.Label(content_main_frame, text=_("Result Content:"), font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
        content_text = tk.Text(content_main_frame, wrap=tk.WORD, height=20, font=("Arial", 11))
        content_text.insert("1.0", selected_report.content)
        content_text.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Scrollbar for content text
        content_scrollbar = ttk.Scrollbar(content_main_frame, orient=tk.VERTICAL, command=content_text.yview)
        content_text.configure(yscrollcommand=content_scrollbar.set)
        content_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Buttons with professional styling
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        def save_result():
            content = content_text.get("1.0", tk.END).strip()
            
            if not content:
                messagebox.showerror(_("Error"), _("Please enter result content"))
                return
            
            # Update the report
            selected_report.content = content
            selected_report.signed_by = self.current_user.id if self.current_user else "N/A"
            selected_report.signed_at = datetime.now()
            
            # Save to database
            try:
                if self.db.update_medical_report(selected_report):
                    messagebox.showinfo(_("Success"), _("Medical result updated successfully"))
                    dialog.destroy()
                    self.load_results_data()
                else:
                    messagebox.showerror(_("Error"), _("Failed to update result"))
            except Exception as e:
                messagebox.showerror(_("Error"), f"{_('Failed to update result')}: {str(e)}")
    
        ttk.Button(button_frame, text=_("💾 Save Result"), command=save_result, style="Accent.TButton").pack(side=tk.RIGHT, padx=5)
        ttk.Button(button_frame, text=_("❌ Cancel"), command=dialog.destroy, style="Accent.TButton").pack(side=tk.RIGHT, padx=5)
        
        # Focus on content text
        content_text.focus()

    def delete_result(self):
        """Delete a selected medical result"""
        selected = self.results_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a result to delete"))
            return
        
        # Confirm deletion
        if not messagebox.askyesno(_("Confirm Delete"), _("Are you sure you want to delete this result?")):
            return
        
        # Get the selected result ID from the item's tags
        item = self.results_tree.selection()[0]
        report_id = self.results_tree.item(item, "tags")[0]
        
        # Delete the report
        try:
            if self.db.delete_medical_report(report_id):
                messagebox.showinfo(_("Success"), _("Result deleted successfully"))
                self.load_results_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to delete result"))
        except Exception as e:
            messagebox.showerror(_("Error"), f"{_('Failed to delete result')}: {str(e)}")

            messagebox.showerror(_("Error"), f"{_('Failed to delete result')}: {str(e)}")

    def print_selected_result(self):
        """Print the selected medical result"""
        selected = self.results_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a result to print"))
            return
        
        # Get the selected result ID from the item's tags
        item = self.results_tree.selection()[0]
        report_id = self.results_tree.item(item, "tags")[0]
        
        # Get report details from database
        selected_report = self.db.get_medical_report(report_id)
        
        if not selected_report:
            messagebox.showerror(_("Error"), _("Result not found"))
            return
        
        # Get associated test request, patient, and test type
        test_request = self.db.get_test_request(selected_report.test_request_id)
        if not test_request:
            messagebox.showerror(_("Error"), _("Test request not found"))
            return
            
        patient = self.db.get_patient(test_request.patient_id)
        if not patient:
            messagebox.showerror(_("Error"), _("Patient not found"))
            return
            
        test_type = self.db.get_test_type(test_request.test_type_id)
        if not test_type:
            messagebox.showerror(_("Error"), _("Test type not found"))
            return
        
        # Create print dialog
        self.create_result_print_dialog(selected_report, patient, test_type, test_request)

    def create_result_print_dialog(self, report, patient, test_type, test_request):
        """Create a print preview dialog for the medical result"""
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Print Medical Result"))
        dialog.geometry("700x800")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Create a text widget for the print content
        text_frame = ttk.Frame(dialog)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create a text widget with scrollbars
        text_widget = tk.Text(text_frame, wrap=tk.WORD, font=("Courier", 10))
        v_scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=text_widget.yview)
        h_scrollbar = ttk.Scrollbar(text_frame, orient=tk.HORIZONTAL, command=text_widget.xview)
        text_widget.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        # Pack the text widget and scrollbars
        text_widget.grid(row=0, column=0, sticky="nsew")
        v_scrollbar.grid(row=0, column=1, sticky="ns")
        h_scrollbar.grid(row=1, column=0, sticky="ew")
        text_frame.grid_rowconfigure(0, weight=1)
        text_frame.grid_columnconfigure(0, weight=1)
        
        # Format the content for printing
        print_content = self.format_result_for_printing(report, patient, test_type, test_request)
        text_widget.insert("1.0", print_content)
        text_widget.config(state=tk.DISABLED)  # Make it read-only
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(button_frame, text=_("Print"), 
                  command=lambda: self.do_print_result(text_widget), 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Save as PDF"), 
                  command=lambda: self.save_result_as_pdf(print_content), 
                  style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Close"), 
                  command=dialog.destroy, 
                  style="Accent.TButton").pack(side=tk.RIGHT, padx=5)

    def format_result_for_printing(self, report, patient, test_type, test_request):
        """Format the result content for printing"""
        # Create a formatted string for the result
        content = []
        content.append("=" * 60)
        content.append(f"{_('MEDICAL LABORATORY RESULT REPORT'):^60}")
        content.append("=" * 60)
        content.append("")
        
        # Report information
        content.append(f"{_('Report ID')}: {report.id}")
        content.append(f"{_('Report Date')}: {report.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        content.append("")
        
        # Patient information
        content.append("-" * 30)
        content.append(f"{_('PATIENT INFORMATION')}")
        content.append("-" * 30)
        content.append(f"{_('Name')}: {patient.name}")
        content.append(f"{_('ID')}: {patient.id}")
        content.append(f"{_('Age')}: {patient.age}")
        content.append(f"{_('Gender')}: {_(patient.gender.value)}")
        content.append(f"{_('Contact')}: {patient.contact_info or _('N/A')}")
        content.append("")
        
        # Test information
        content.append("-" * 30)
        content.append(f"{_('TEST INFORMATION')}")
        content.append("-" * 30)
        content.append(f"{_('Test Name')}: {test_type.name}")
        content.append(f"{_('Test Category')}: {test_type.category}")
        content.append(f"{_('Test ID')}: {test_request.id}")
        content.append(f"{_('Requested By')}: {test_request.requested_by}")
        content.append(f"{_('Requested At')}: {test_request.requested_at.strftime('%Y-%m-%d %H:%M')}")
        content.append(f"{_('Status')}: {_(test_request.status.value)}")
        content.append("")
        
        # Result content
        content.append("-" * 30)
        content.append(f"{_('RESULT DETAILS')}")
        content.append("-" * 30)
        content.append(report.content)
        content.append("")
        
        # Signature information
        content.append("-" * 30)
        content.append(f"{_('SIGNATURE INFORMATION')}")
        content.append("-" * 30)
        content.append(f"{_('Signed By')}: {report.signed_by if report.signed_by != 'N/A' else _('Not signed yet')}")
        content.append(f"{_('Signed At')}: {report.signed_at.strftime('%Y-%m-%d %H:%M') if report.signed_at else _('Not signed yet')}")
        content.append("")
        
        # Footer
        content.append("=" * 60)
        content.append(f"{_('Generated by Medical Laboratory Management System')}")
        content.append(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content.append("=" * 60)
        
        return "\n".join(content)

    def do_print_result(self, text_widget):
        """Actually print the result"""
        try:
            # Get the content from the text widget
            content = text_widget.get("1.0", tk.END)
            
            # In a real implementation, we would use the system's print dialog
            # For now, we'll show a message indicating what would happen
            messagebox.showinfo(
                _("Print"), 
                _("In a full implementation, this would send the following content to your printer:\n\n") + 
                content[:200] + "..."
            )
        except Exception as e:
            messagebox.showerror(_("Error"), f"{_('Failed to print result')}: {str(e)}")

    def save_result_as_pdf(self, content):
        """Save the result as a PDF file"""
        try:
            # Ask user for file location
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[(_("PDF files"), "*.pdf"), (_("All files"), "*.*")],
                title=_("Save Result as PDF")
            )
            
            if file_path:
                # In a real implementation, we would convert the content to PDF
                # For now, we'll save as a text file and show a message
                with open(file_path.replace('.pdf', '.txt'), 'w') as f:
                    f.write(content)
                
                messagebox.showinfo(
                    _("Save as PDF"), 
                    _("In a full implementation, this would save the result as a PDF file.\n\n") +
                    _("For now, a text file has been saved at: {}").format(file_path.replace('.pdf', '.txt'))
                )
        except Exception as e:
            messagebox.showerror(_("Error"), f"{_('Failed to save result as PDF')}: {str(e)}")

    def print_invoice(self):
        """Print the selected invoice"""
        selected = self.billing_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select an invoice to print"))
            return
    
        # Get the selected invoice data
        item = self.billing_tree.selection()[0]
        values = self.billing_tree.item(item, "values")
    
        invoice_id = values[0]
        patient_name = values[1]
        amount = values[2]
        paid = values[3]
        status = values[4]
        date = values[5]
    
        # Create print dialog
        self.create_invoice_print_dialog(invoice_id, patient_name, amount, paid, status, date)

    def create_invoice_print_dialog(self, invoice_id, patient_name, amount, paid, status, date):
        """Create a print preview dialog for the invoice"""
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Print Invoice"))
        dialog.geometry("700x800")
        dialog.transient(self.root)
        dialog.grab_set()
    
        # Create a text widget for the print content
        text_frame = ttk.Frame(dialog)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    
        # Create a text widget with scrollbars
        text_widget = tk.Text(text_frame, wrap=tk.WORD, font=("Courier", 10))
        v_scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=text_widget.yview)
        h_scrollbar = ttk.Scrollbar(text_frame, orient=tk.HORIZONTAL, command=text_widget.xview)
        text_widget.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
    
        # Pack the text widget and scrollbars
        text_widget.grid(row=0, column=0, sticky="nsew")
        v_scrollbar.grid(row=0, column=1, sticky="ns")
        h_scrollbar.grid(row=1, column=0, sticky="ew")
        text_frame.grid_rowconfigure(0, weight=1)
        text_frame.grid_columnconfigure(0, weight=1)
    
        # Format the content for printing
        print_content = self.format_invoice_for_printing(invoice_id, patient_name, amount, paid, status, date)
        text_widget.insert("1.0", print_content)
    
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
    
        ttk.Button(button_frame, text=_("Print"), 
                  command=lambda: self.do_print_invoice(text_widget)).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Close"), 
                  command=dialog.destroy).pack(side=tk.RIGHT, padx=5)

    def format_invoice_for_printing(self, invoice_id, patient_name, amount, paid, status, date):
        """Format the invoice content for printing"""
        # Create a formatted string for the invoice
        content = []
        content.append("=" * 60)
        content.append(f"{_('MEDICAL LABORATORY INVOICE'):^60}")
        content.append("=" * 60)
        content.append("")
    
        # Invoice information
        content.append(f"{_('Invoice ID')}: {invoice_id}")
        content.append(f"{_('Invoice Date')}: {date}")
        content.append("")
    
        # Patient information
        content.append("-" * 30)
        content.append(f"{_('PATIENT INFORMATION')}")
        content.append("-" * 30)
        content.append(f"{_('Patient Name')}: {patient_name}")
        content.append("")
    
        # Billing information
        content.append("-" * 30)
        content.append(f"{_('BILLING INFORMATION')}")
        content.append("-" * 30)
        content.append(f"{_('Total Amount')}: {amount}")
        content.append(f"{_('Paid Amount')}: {paid}")
        content.append(f"{_('Status')}: {status}")
        content.append("")
    
        # Payment details (in a real app, this would include individual tests)
        content.append("-" * 30)
        content.append(f"{_('PAYMENT DETAILS')}")
        content.append("-" * 30)
        content.append(f"{_('Test'):<30} {_('Price')}")
        content.append("-" * 30)
        content.append(f"{_('Complete Blood Count'):<30} {'$50.00'}")
        content.append(f"{_('Urinalysis'):<30} {'$30.00'}")
        content.append(f"{_('Stool Analysis'):<30} {'$40.00'}")
        content.append("-" * 30)
        content.append(f"{'Total':<30} {amount}")
        content.append("")
    
        # Footer
        content.append("=" * 60)
        content.append(f"{_('Generated by Medical Laboratory Management System')}")
        content.append(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content.append("=" * 60)
    
        return "\n".join(content)

    def do_print_invoice(self, text_widget):
        """Actually print the invoice"""
        try:
            # Get the content from the text widget
            content = text_widget.get("1.0", tk.END)
            
            # In a real implementation, we would use the system's print dialog
            # For now, we'll show a message indicating what would happen
            messagebox.showinfo(
                _("Print"), 
                _("In a full implementation, this would send the following content to your printer:\n\n") + 
                content[:200] + "..."
            )
        except Exception as e:
            messagebox.showerror(_("Error"), f"{_('Failed to print invoice')}: {str(e)}")

    def manage_test_templates(self):
        """Manage test templates for different test types with professional UI"""
        # Create professional template management dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("📄 Professional Test Template Management"))
        dialog.geometry("900x700")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Professional header
        header_frame = ttk.Frame(dialog, style="Header.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("📄 Test Template Management System"), 
                 font=("Arial", 16, "bold"), style="Header.TLabel").pack()

    
        # Create notebook for different sections with enhanced styling
        notebook = ttk.Notebook(dialog)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    
        # Template management tab
        manage_frame = ttk.Frame(notebook)
        notebook.add(manage_frame, text=_("🛠️ Manage Templates"))
    
        # Test type selection with enhanced UI
        test_frame = ttk.LabelFrame(manage_frame, text=_("🧪 Select Test Type"), padding=15)
        test_frame.pack(fill=tk.X, padx=15, pady=10)
    
        ttk.Label(test_frame, text=_("Test Type:"), font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
        test_type_var = tk.StringVar()
        test_type_combo = ttk.Combobox(test_frame, textvariable=test_type_var, state="readonly", width=60, font=("Arial", 10))
        test_type_combo.pack(fill=tk.X, pady=5)
    
        # Load test types
        test_types = self.db.get_all_test_types()
        test_type_map = {t.name: t for t in test_types}
        test_type_combo['values'] = [t.name for t in test_types]
    
        # Template content with enhanced styling
        content_frame = ttk.LabelFrame(manage_frame, text=_("📝 Template Content"), padding=15)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)
    
        # Create text widget with scrollbars and enhanced styling
        text_frame = ttk.Frame(content_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)
    
        template_text = tk.Text(text_frame, wrap=tk.WORD, font=("Arial", 11))
        v_scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=template_text.yview)
        h_scrollbar = ttk.Scrollbar(text_frame, orient=tk.HORIZONTAL, command=template_text.xview)
        template_text.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
    
        # Pack the text widget and scrollbars
        template_text.grid(row=0, column=0, sticky="nsew")
        v_scrollbar.grid(row=0, column=1, sticky="ns")
        h_scrollbar.grid(row=1, column=0, sticky="ew")
        text_frame.grid_rowconfigure(0, weight=1)
        text_frame.grid_columnconfigure(0, weight=1)
    
        # Template buttons with enhanced styling
        button_frame = ttk.Frame(manage_frame)
        button_frame.pack(fill=tk.X, padx=15, pady=10)
    
        # Load from Word file button
        def load_from_word():
            try:
                file_path = filedialog.askopenfilename(
                    title=_("Select Word Template File"),
                    filetypes=[(_("Word files"), "*.docx"), (_("All files"), "*.*")]
                )
                
                if file_path:
                    # Load content from Word document
                    doc = Document(file_path)
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    template_text.delete("1.0", tk.END)
                    template_text.insert("1.0", content)
                    messagebox.showinfo(_("Success"), _("Template loaded successfully from Word file"))
            except Exception as e:
                messagebox.showerror(_("Error"), f"{_('Failed to load template')}: {str(e)}")
    
        # Save template button
        def save_template():
            test_type_name = test_type_var.get()
            content = template_text.get("1.0", tk.END).strip()
            
            if not test_type_name:
                messagebox.showwarning(_("Warning"), _("Please select a test type"))
                return
            
            if not content:
                messagebox.showwarning(_("Warning"), _("Please enter template content"))
                return
            
            # Get test type
            test_type = test_type_map.get(test_type_name)
            if not test_type:
                messagebox.showerror(_("Error"), _("Invalid test type selection"))
                return
            
            # Check if template already exists for this test type
            existing_template = self.db.get_test_template_by_test_type(test_type.id)
            
            if existing_template:
                # Update existing template
                existing_template.template_content = content
                existing_template.updated_at = datetime.now()
                if self.db.update_test_template(existing_template):
                    messagebox.showinfo(_("Success"), _("Template updated successfully for {}").format(test_type_name))
                else:
                    messagebox.showerror(_("Error"), _("Failed to update template"))
            else:
                # Create new template
                new_template = TestTemplate(
                    id=str(uuid.uuid4()),
                    test_type_id=test_type.id,
                    template_content=content
                )
                if self.db.create_test_template(new_template):
                    messagebox.showinfo(_("Success"), _("Template saved successfully for {}").format(test_type_name))
                else:
                    messagebox.showerror(_("Error"), _("Failed to save template"))
    
        # Load template button
        def load_template():
            test_type_name = test_type_var.get()
            
            if not test_type_name:
                messagebox.showwarning(_("Warning"), _("Please select a test type"))
                return
            
            # Get test type
            test_type = test_type_map.get(test_type_name)
            if not test_type:
                messagebox.showerror(_("Error"), _("Invalid test type selection"))
                return
            
            # Get template for this test type
            template = self.db.get_test_template_by_test_type(test_type.id)
            
            if template:
                template_text.delete("1.0", tk.END)
                template_text.insert("1.0", template.template_content)
                messagebox.showinfo(_("Success"), _("Template loaded successfully for {}").format(test_type_name))
            else:
                template_text.delete("1.0", tk.END)
                messagebox.showinfo(_("Info"), _("No template found for {}. You can create one now.").format(test_type_name))
    
        # Enhanced buttons with icons
        ttk.Button(button_frame, text=_("📂 Load from Word File"), 
                  command=load_from_word, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("📥 Load Template"), 
                  command=load_template, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("💾 Save Template"), 
                  command=save_template, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
    
        # Template preview tab
        preview_frame = ttk.Frame(notebook)
        notebook.add(preview_frame, text=_("👁️ Preview Templates"))
    
        # Preview controls with enhanced styling
        preview_controls = ttk.Frame(preview_frame)
        preview_controls.pack(fill=tk.X, padx=15, pady=10)
    
        ttk.Label(preview_controls, text=_("Select Test Type for Preview:"), font=("Arial", 10, "bold")).pack(side=tk.LEFT)
        preview_test_var = tk.StringVar()
        preview_test_combo = ttk.Combobox(preview_controls, textvariable=preview_test_var, 
                                         state="readonly", width=35, font=("Arial", 10))
        preview_test_combo.pack(side=tk.LEFT, padx=5)
        preview_test_combo['values'] = [t.name for t in test_types]
    
        # Preview button with enhanced styling
        def preview_template():
            test_type_name = preview_test_var.get()
            
            if not test_type_name:
                messagebox.showwarning(_("Warning"), _("Please select a test type"))
                return
            
            # Get test type
            test_type = test_type_map.get(test_type_name)
            if not test_type:
                messagebox.showerror(_("Error"), _("Invalid test type selection"))
                return
            
            # Get template for this test type
            template = self.db.get_test_template_by_test_type(test_type.id)
            
            if template:
                preview_text.config(state=tk.NORMAL)
                preview_text.delete("1.0", tk.END)
                preview_text.insert("1.0", template.template_content)
                preview_text.config(state=tk.DISABLED)
                messagebox.showinfo(_("Success"), _("Template preview loaded for {}").format(test_type_name))
            else:
                preview_text.config(state=tk.NORMAL)
                preview_text.delete("1.0", tk.END)
                preview_text.insert("1.0", _("No template found for {}").format(test_type_name))
                preview_text.config(state=tk.DISABLED)
    
        ttk.Button(preview_controls, text=_("👁️ Preview Template"), 
                  command=preview_template, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
    
        # Preview content with enhanced styling
        preview_content_frame = ttk.LabelFrame(preview_frame, text=_("📄 Template Preview"), padding=15)
        preview_content_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)
    
        # Create text widget with scrollbars for preview
        preview_text_frame = ttk.Frame(preview_content_frame)
        preview_text_frame.pack(fill=tk.BOTH, expand=True)
    
        preview_text = tk.Text(preview_text_frame, wrap=tk.WORD, state=tk.DISABLED, font=("Arial", 11))
        preview_v_scrollbar = ttk.Scrollbar(preview_text_frame, orient=tk.VERTICAL, command=preview_text.yview)
        preview_h_scrollbar = ttk.Scrollbar(preview_text_frame, orient=tk.HORIZONTAL, command=preview_text.xview)
        preview_text.configure(yscrollcommand=preview_v_scrollbar.set, xscrollcommand=preview_h_scrollbar.set)
    
        # Pack the text widget and scrollbars
        preview_text.grid(row=0, column=0, sticky="nsew")
        preview_v_scrollbar.grid(row=0, column=1, sticky="ns")
        preview_h_scrollbar.grid(row=1, column=0, sticky="ew")
        preview_text_frame.grid_rowconfigure(0, weight=1)
        preview_text_frame.grid_columnconfigure(0, weight=1)
    
        # Close button with enhanced styling
        close_frame = ttk.Frame(dialog)
        close_frame.pack(fill=tk.X, padx=10, pady=10)
    
        ttk.Button(close_frame, text=_("❌ Close"), 
                  command=dialog.destroy, style="Accent.TButton").pack(side=tk.RIGHT, padx=5)
    
    def show_reports(self):
        self.current_screen = self.show_reports
        self.clear_content()
        
        # Reports header with 3D styling
        header_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("Medical Reports"), 
                 style="Title.TLabel").pack(side=tk.LEFT)
        
        ttk.Button(header_frame, text=_("Create Report"), 
                  command=self.create_report, style="Accent.TButton").pack(side=tk.RIGHT)
        
        # Reports table with enhanced styling
        table_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        table_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create treeview with custom styling
        columns = (_("ID"), _("Test Request"), _("Signed By"), _("Signed At"), _("Status"))
        self.reports_tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        
        for col in columns:
            self.reports_tree.heading(col, text=col)
            self.reports_tree.column(col, width=120)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, 
                                 command=self.reports_tree.yview)
        self.reports_tree.configure(yscroll=scrollbar.set)
        
        self.reports_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load reports data
        self.load_reports_data()
        
        # Action buttons with 3D styling
        action_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        action_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(action_frame, text=_("View Report"), 
                  command=self.view_report, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Sign Report"), 
                  command=self.sign_report, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Send Report"), 
                  command=self.send_report, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
    
    def load_reports_data(self):
        # Check if reports_tree exists
        if not hasattr(self, 'reports_tree'):
            return
            
        # Clear existing data
        for item in self.reports_tree.get_children():
            self.reports_tree.delete(item)
        
        # Load medical reports from database
        reports = self.db.get_all_medical_reports()
        
        for report in reports:
            # Get test request to get patient and test type info
            test_request = self.db.get_test_request(report.test_request_id)
            patient_name = "Unknown Patient"
            test_name = "Unknown Test"
            
            if test_request:
                # Get patient info
                patient = self.db.get_patient(test_request.patient_id)
                if patient:
                    patient_name = patient.name
                
                # Get test type info
                test_type = self.db.get_test_type(test_request.test_type_id)
                if test_type:
                    test_name = test_type.name
            
            self.reports_tree.insert("", tk.END, values=(
                report.id[:8],  # Short ID for display
                f"{patient_name} - {test_name}",
                report.signed_by if report.signed_by != "N/A" else _("Not signed"),
                report.signed_at.strftime("%Y-%m-%d %H:%M") if report.signed_at else _("Not signed"),
                _("Signed") if report.signed_by != "N/A" else _("Pending")
            ))
    
    def load_results_data(self):
        # Check if results_tree exists
        if not hasattr(self, 'results_tree'):
            return
            
        # Clear existing data
        for item in self.results_tree.get_children():
            self.results_tree.delete(item)
        
        # Load medical reports from database
        reports = self.db.get_all_medical_reports()
        
        for report in reports:
            # Get test request to get patient and test type info
            test_request = self.db.get_test_request(report.test_request_id)
            patient_name = _("Unknown Patient")
            test_name = _("Unknown Test")
            status = _("Pending")
            
            if test_request:
                # Get patient info
                patient = self.db.get_patient(test_request.patient_id)
                if patient:
                    patient_name = patient.name
                
                # Get test type info
                test_type = self.db.get_test_type(test_request.test_type_id)
                if test_type:
                    test_name = test_type.name
                
                # Set status
                status = _(test_request.status.value)
            
            # Insert item and store the full ID in the item's values
            item_id = self.results_tree.insert("", tk.END, values=(
                report.id[:8],  # Short ID for display
                patient_name,
                test_name,
                status,
                report.created_at.strftime("%Y-%m-%d %H:%M")
            ))
            # Store the full ID in the item's tags for later retrieval
            self.results_tree.item(item_id, tags=(report.id,))
            # Store the full ID in the item's tags for later retrieval
            self.results_tree.item(item_id, tags=(report.id,))
    
    def create_report(self):
        # Create report dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Create Medical Report"))
        dialog.geometry("500x400")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Form fields
        ttk.Label(dialog, text=_("Test Request ID:")).pack(pady=5)
        request_entry = ttk.Entry(dialog, width=50)
        request_entry.pack(pady=5)
        
        ttk.Label(dialog, text=_("Report Content:")).pack(pady=5)
        content_text = tk.Text(dialog, width=50, height=15)
        content_text.pack(pady=5)
        
        def save_report():
            request_id = request_entry.get().strip()
            content = content_text.get("1.0", tk.END).strip()
            
            # Validation
            if not request_id:
                messagebox.showerror(_("Error"), _("Please enter test request ID"))
                return
            
            if not content:
                messagebox.showerror(_("Error"), _("Please enter report content"))
                return
            
            # Create report
            report = MedicalReport(
                id=str(uuid.uuid4()),
                test_request_id=request_id,
                content=content,
                signed_by=self.current_user.id if self.current_user else "N/A",
                signed_at=datetime.now()
            )
            
            if self.db.create_medical_report(report):
                messagebox.showinfo(_("Success"), _("Report created successfully"))
                dialog.destroy()
                self.load_reports_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to create report"))
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text=_("Save"), command=save_report).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), 
                  command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Focus on first entry
        request_entry.focus()
    
    def view_report(self):
        selected = self.reports_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a report"))
            return
        
        # In a real app, this would show the report details
        messagebox.showinfo(_("View Report"), _("Report viewing functionality would be implemented here"))
    
    def sign_report(self):
        selected = self.reports_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a report"))
            return
        
        # In a real app, this would sign the report
        messagebox.showinfo(_("Sign Report"), _("Report signing functionality would be implemented here"))
    
    def send_report(self):
        selected = self.reports_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a report"))
            return
        
        # In a real app, this would send the report
        messagebox.showinfo(_("Send Report"), _("Report sending functionality would be implemented here"))
    
    def show_billing(self):
        self.current_screen = self.show_billing
        self.clear_content()
        
        # Billing header with 3D styling
        header_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("Billing and Payments"), 
                 style="Title.TLabel").pack(side=tk.LEFT)
        
        ttk.Button(header_frame, text=_("Create Invoice"), 
                  command=self.create_invoice, style="Accent.TButton").pack(side=tk.RIGHT)
        
        # Billing table with enhanced styling
        table_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        table_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create treeview with custom styling
        columns = (_("ID"), _("Patient"), _("Amount"), _("Paid"), _("Status"), _("Date"))
        self.billing_tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        
        for col in columns:
            self.billing_tree.heading(col, text=col)
            self.billing_tree.column(col, width=100)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, 
                                 command=self.billing_tree.yview)
        self.billing_tree.configure(yscroll=scrollbar.set)
        
        self.billing_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load billing data
        self.load_billing_data()
        
        # Action buttons with 3D styling
        action_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        action_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(action_frame, text=_("View Invoice"), 
                  command=self.view_invoice, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Process Payment"), 
                  command=self.process_payment, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Delete Invoice"), 
                  command=self.delete_invoice, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Generate Report"), 
                  command=self.generate_billing_report, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
    
    def load_billing_data(self):
        # Check if billing_tree exists
        if not hasattr(self, 'billing_tree'):
            return
            
        # Clear existing data
        for item in self.billing_tree.get_children():
            self.billing_tree.delete(item)
        
        # In a real app, this would load billing data from database
        # For now, we'll show sample data
        sample_invoices = [
            ("INV-001", "John Doe", "$150.00", "$150.00", _("Paid"), "2023-05-15"),
            ("INV-002", "Jane Smith", "$85.50", "$50.00", _("Partial"), "2023-05-16"),
            ("INV-003", "Robert Johnson", "$200.00", "$0.00", _("Unpaid"), "2023-05-17")
        ]
        
        for invoice in sample_invoices:
            self.billing_tree.insert("", tk.END, values=invoice)
    
    def create_invoice(self):
        # Create invoice dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Create Invoice"))
        dialog.geometry("600x500")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Variables to store selected tests and their prices
        self.selected_tests = []
        self.test_prices = {}
        
        # Patient information section
        patient_frame = ttk.LabelFrame(dialog, text=_("Patient Information"), padding=10)
        patient_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Patient name
        ttk.Label(patient_frame, text=_("Patient Name:")).grid(row=0, column=0, sticky=tk.W, pady=5)
        patient_name_entry = ttk.Entry(patient_frame, width=30)
        patient_name_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        
        # Patient ID/Number
        ttk.Label(patient_frame, text=_("Patient Number:")).grid(row=1, column=0, sticky=tk.W, pady=5)
        patient_id_entry = ttk.Entry(patient_frame, width=30)
        patient_id_entry.grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)
        
        # Tests selection section
        tests_frame = ttk.LabelFrame(dialog, text=_("Tests Performed"), padding=10)
        tests_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Available tests listbox
        ttk.Label(tests_frame, text=_("Available Tests:")).grid(row=0, column=0, sticky=tk.W, pady=5)
        available_frame = ttk.Frame(tests_frame)
        available_frame.grid(row=1, column=0, padx=5, pady=5, sticky="nsew")
        tests_frame.grid_rowconfigure(1, weight=1)
        tests_frame.grid_columnconfigure(0, weight=1)
        
        available_listbox = tk.Listbox(available_frame, selectmode=tk.MULTIPLE, height=6)
        available_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        available_scrollbar = ttk.Scrollbar(available_frame, orient=tk.VERTICAL, command=available_listbox.yview)
        available_listbox.configure(yscrollcommand=available_scrollbar.set)
        available_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Add selected test button
        add_test_btn = ttk.Button(tests_frame, text=_("Add Selected Tests"))
        add_test_btn.grid(row=2, column=0, pady=5)
        
        # Selected tests listbox with prices
        ttk.Label(tests_frame, text=_("Selected Tests:")).grid(row=3, column=0, sticky=tk.W, pady=(10, 5))
        selected_frame = ttk.Frame(tests_frame)
        selected_frame.grid(row=4, column=0, padx=5, pady=5, sticky="nsew")
        tests_frame.grid_rowconfigure(4, weight=1)
        
        # Create treeview for selected tests with test name and price columns
        selected_columns = (_("Test Name"), _("Price"))
        selected_tree = ttk.Treeview(selected_frame, columns=selected_columns, show="headings", height=6)
        selected_tree.heading(_("Test Name"), text=_("Test Name"))
        selected_tree.heading(_("Price"), text=_("Price"))
        selected_tree.column(_("Test Name"), width=200)
        selected_tree.column(_("Price"), width=100)
        
        selected_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        selected_scrollbar = ttk.Scrollbar(selected_frame, orient=tk.VERTICAL, command=selected_tree.yview)
        selected_tree.configure(yscrollcommand=selected_scrollbar.set)
        selected_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Remove test button
        remove_test_btn = ttk.Button(tests_frame, text=_("Remove Selected Test"))
        remove_test_btn.grid(row=5, column=0, pady=5)
        
        # Total amount section
        total_frame = ttk.Frame(dialog)
        total_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(total_frame, text=_("Total Amount:"), font=("Arial", 12, "bold")).pack(side=tk.LEFT, padx=5)
        total_amount_var = tk.StringVar(value="$0.00")
        total_amount_label = ttk.Label(total_frame, textvariable=total_amount_var, font=("Arial", 12, "bold"))
        total_amount_label.pack(side=tk.RIGHT, padx=5)
        
        # Load available tests from database
        test_types = self.db.get_all_test_types()
        for test in test_types:
            display_text = f"{test.name} - ${test.price:.2f}"
            available_listbox.insert(tk.END, display_text)
            # Store test info for later use
            self.test_prices[display_text] = test.price
    
        def add_selected_tests():
            """Add selected tests to the selected tests list"""
            selected_indices = available_listbox.curselection()
            for index in selected_indices:
                test_display = available_listbox.get(index)
                if test_display not in self.selected_tests:
                    self.selected_tests.append(test_display)
                    price = self.test_prices.get(test_display, 0.0)
                    selected_tree.insert("", tk.END, values=(test_display, f"${price:.2f}"))
            update_total()
    
        def remove_selected_test():
            """Remove selected test from the selected tests list"""
            selected_items = selected_tree.selection()
            for item in selected_items:
                values = selected_tree.item(item, "values")
                test_display = values[0]
                if test_display in self.selected_tests:
                    self.selected_tests.remove(test_display)
                selected_tree.delete(item)
            update_total()
    
        def update_total():
            """Calculate and update the total amount"""
            total = 0.0
            for test_display in self.selected_tests:
                total += self.test_prices.get(test_display, 0.0)
            total_amount_var.set(f"${total:.2f}")
    
        def save_invoice():
            patient_name = patient_name_entry.get().strip()
            patient_id = patient_id_entry.get().strip()
            
            # Validation
            if not patient_name:
                messagebox.showerror(_("Error"), _("Please enter patient name"))
                return
            
            if not patient_id:
                messagebox.showerror(_("Error"), _("Please enter patient number"))
                return
            
            if not self.selected_tests:
                messagebox.showerror(_("Error"), _("Please select at least one test"))
                return
            
            # Get total amount
            total_amount_str = total_amount_var.get().replace("$", "")
            try:
                total_amount = float(total_amount_str)
            except ValueError:
                messagebox.showerror(_("Error"), _("Invalid total amount"))
                return
            
            # For demo purposes, we'll create a simple list of test request IDs
            # In a real application, you would create actual test requests
            request_ids = [f"REQ-{i+1}" for i in range(len(self.selected_tests))]
            
            # Create invoice
            invoice = Invoice(
                id=str(uuid.uuid4()),
                patient_id=patient_id,
                test_request_ids=request_ids,
                total_amount=total_amount,
                paid_amount=0.0
            )
            
            # In a real app, this would save to database
            messagebox.showinfo(_("Success"), _("Invoice created successfully"))
            dialog.destroy()
            self.load_billing_data()
    
        # Bind buttons to functions
        add_test_btn.config(command=add_selected_tests)
        remove_test_btn.config(command=remove_selected_test)
    
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=20)
        
        ttk.Button(button_frame, text=_("Save Invoice"), command=save_invoice).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), command=dialog.destroy).pack(side=tk.RIGHT, padx=5)
        
        # Focus on first entry
        patient_name_entry.focus()

    def view_invoice(self):
        messagebox.showinfo(_("View Invoice"), _("Invoice viewing functionality would be implemented here"))
    
    def process_payment(self):
        messagebox.showinfo(_("Process Payment"), _("Payment processing functionality would be implemented here"))
    
    def generate_billing_report(self):
        messagebox.showinfo(_("Generate Report"), _("Billing report generation functionality would be implemented here"))
    
    def delete_invoice(self):
        """Delete selected invoice"""
        selected = self.billing_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select an invoice to delete"))
            return
    
        # Confirm deletion
        result = messagebox.askyesno(_("Confirm Deletion"), _("Are you sure you want to delete this invoice?"))
        if result:
            # In a real app, this would delete from database
            messagebox.showinfo(_("Success"), _("Invoice deleted successfully"))
            self.load_billing_data()
    
    def show_inventory(self):
        self.current_screen = self.show_inventory
        self.clear_content()
        
        # Inventory header with 3D styling
        header_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("Inventory Management"), 
                 style="Title.TLabel").pack(side=tk.LEFT)
        
        ttk.Button(header_frame, text=_("Add Item"), 
                  command=self.add_inventory_item, style="Accent.TButton").pack(side=tk.RIGHT)
        
        # Inventory table with enhanced styling
        table_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        table_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create treeview with custom styling
        columns = (_("ID"), _("Name"), _("Quantity"), _("Min Quantity"), _("Supplier"), _("Expiry Date"))
        self.inventory_tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        
        for col in columns:
            self.inventory_tree.heading(col, text=col)
            self.inventory_tree.column(col, width=100)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, 
                                 command=self.inventory_tree.yview)
        self.inventory_tree.configure(yscroll=scrollbar.set)
        
        self.inventory_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load inventory data
        self.load_inventory_data()
        
        # Action buttons with 3D styling
        action_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        action_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(action_frame, text=_("View Details"), 
                  command=self.view_inventory_item, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Update Quantity"), 
                  command=self.update_inventory_quantity, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Create Purchase Order"), 
                  command=self.create_purchase_order, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Low Stock Alert"), 
                  command=self.show_low_stock_alerts, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
    
    def load_inventory_data(self):
        # Clear existing data
        for item in self.inventory_tree.get_children():
            self.inventory_tree.delete(item)
        
        # In a real app, this would load inventory items from database
        # For now, we'll show a message
        self.inventory_tree.insert("", tk.END, values=(
            _("N/A"), _("N/A"), _("N/A"), _("N/A"), _("N/A"), _("N/A")
        ))
    
    def add_inventory_item(self):
        # Create add inventory item dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Add Inventory Item"))
        dialog.geometry("400x400")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Form fields
        ttk.Label(dialog, text=_("Item Name:")).pack(pady=5)
        name_entry = ttk.Entry(dialog, width=40)
        name_entry.pack(pady=5)
        
        ttk.Label(dialog, text=_("Description:")).pack(pady=5)
        desc_text = tk.Text(dialog, width=40, height=3)
        desc_text.pack(pady=5)
        
        ttk.Label(dialog, text=_("Quantity:")).pack(pady=5)
        quantity_entry = ttk.Entry(dialog, width=40)
        quantity_entry.pack(pady=5)
        
        ttk.Label(dialog, text=_("Minimum Quantity:")).pack(pady=5)
        min_quantity_entry = ttk.Entry(dialog, width=40)
        min_quantity_entry.pack(pady=5)
        
        ttk.Label(dialog, text=_("Supplier:")).pack(pady=5)
        supplier_entry = ttk.Entry(dialog, width=40)
        supplier_entry.pack(pady=5)
        
        ttk.Label(dialog, text=_("Expiry Date (YYYY-MM-DD):")).pack(pady=5)
        expiry_entry = ttk.Entry(dialog, width=40)
        expiry_entry.pack(pady=5)
        
        def save_inventory_item():
            name = name_entry.get().strip()
            description = desc_text.get("1.0", tk.END).strip()
            quantity_str = quantity_entry.get().strip()
            min_quantity_str = min_quantity_entry.get().strip()
            supplier = supplier_entry.get().strip()
            expiry_str = expiry_entry.get().strip()
            
            # Validation
            if not name:
                messagebox.showerror(_("Error"), _("Please enter item name"))
                return
            
            try:
                quantity = int(quantity_str)
                if quantity < 0:
                    raise ValueError()
            except ValueError:
                messagebox.showerror(_("Error"), _("Please enter a valid quantity"))
                return
            
            try:
                min_quantity = int(min_quantity_str)
                if min_quantity < 0:
                    raise ValueError()
            except ValueError:
                messagebox.showerror(_("Error"), _("Please enter a valid minimum quantity"))
                return
            
            # Parse expiry date if provided
            expiry_date = None
            if expiry_str:
                try:
                    expiry_date = datetime.strptime(expiry_str, "%Y-%m-%d")
                except ValueError:
                    messagebox.showerror(_("Error"), _("Please enter a valid expiry date (YYYY-MM-DD)"))
                    return
            
            # Create inventory item
            item = InventoryItem(
                id=str(uuid.uuid4()),
                name=name,
                description=description,
                quantity=quantity,
                min_quantity=min_quantity,
                supplier=supplier,
                expiry_date=expiry_date
            )
            
            # In a real app, this would save to database
            messagebox.showinfo(_("Success"), _("Inventory item added successfully"))
            dialog.destroy()
            self.load_inventory_data()
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text=_("Save"), command=save_inventory_item).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), 
                  command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Focus on first entry
        name_entry.focus()
    
    def view_inventory_item(self):
        messagebox.showinfo(_("View Item"), _("Inventory item details would be shown here"))
    
    def update_inventory_quantity(self):
        messagebox.showinfo(_("Update Quantity"), _("Inventory quantity update functionality would be implemented here"))
    
    def create_purchase_order(self):
        messagebox.showinfo(_("Purchase Order"), _("Purchase order creation functionality would be implemented here"))
    
    def show_low_stock_alerts(self):
        messagebox.showinfo(_("Low Stock Alerts"), _("Low stock alerts would be shown here"))
    
    def show_results(self):
        self.current_screen = self.show_results
        self.clear_content()
        
        # Results header with 3D styling
        header_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("Results Management"), 
                 style="Title.TLabel").pack(side=tk.LEFT)
        
        ttk.Button(header_frame, text=_("Add Result"), 
                  command=self.create_new_result, style="Accent.TButton").pack(side=tk.RIGHT)
        
        # Results table with enhanced styling
        table_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        table_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create treeview with custom styling
        columns = (_("ID"), _("Patient"), _("Tests"), _("Signed By"), _("Signed At"))
        self.results_tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        
        for col in columns:
            self.results_tree.heading(col, text=col)
            self.results_tree.column(col, width=120)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, 
                                 command=self.results_tree.yview)
        self.results_tree.configure(yscroll=scrollbar.set)
        
        self.results_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load results data
        self.load_results_data()
        
        # Action buttons with 3D styling
        action_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        action_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(action_frame, text=_("View Details"), 
                  command=self.view_result_details, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Edit Result"), 
                  command=self.edit_result, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Delete Result"), 
                  command=self.delete_result, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Print Result"), 
                  command=self.print_selected_result, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
    
    def show_billing(self):
        self.current_screen = self.show_billing
        self.clear_content()
        
        # Billing header with 3D styling
        header_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("Billing and Payments"), 
             style="Title.TLabel").pack(side=tk.LEFT)
        
        ttk.Button(header_frame, text=_("Create Invoice"), 
              command=self.create_invoice, style="Accent.TButton").pack(side=tk.RIGHT)
        
        # Billing table with enhanced styling
        table_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        table_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create treeview with custom styling
        columns = (_("ID"), _("Patient"), _("Amount"), _("Paid"), _("Status"), _("Date"))
        self.billing_tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        
        for col in columns:
            self.billing_tree.heading(col, text=col)
            self.billing_tree.column(col, width=100)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, 
                             command=self.billing_tree.yview)
        self.billing_tree.configure(yscroll=scrollbar.set)
        
        self.billing_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load billing data
        self.load_billing_data()
        
        # Action buttons with 3D styling
        action_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        action_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(action_frame, text=_("View Invoice"), 
              command=self.view_invoice, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Process Payment"), 
              command=self.process_payment, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Print Invoice"), 
              command=self.print_invoice, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Delete Invoice"), 
              command=self.delete_invoice, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Generate Report"), 
              command=self.generate_billing_report, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
    
    def show_users(self):
        self.current_screen = self.show_users
        self.clear_content()
        
        # Users header with 3D styling
        header_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(header_frame, text=_("User Management"), 
                 style="Title.TLabel").pack(side=tk.LEFT)
        
        ttk.Button(header_frame, text=_("Add User"), 
                  command=self.add_user, style="Accent.TButton").pack(side=tk.RIGHT)
        
        # Users table with enhanced styling
        table_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        table_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create treeview with custom styling
        columns = (_("ID"), _("Username"), _("Email"), _("Role"), _("Active"), _("Last Login"))
        self.users_tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        
        for col in columns:
            self.users_tree.heading(col, text=col)
            self.users_tree.column(col, width=100)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, 
                                 command=self.users_tree.yview)
        self.users_tree.configure(yscroll=scrollbar.set)
        
        self.users_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Load users data
        self.load_users_data()
        
        # Action buttons with 3D styling
        action_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        action_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(action_frame, text=_("View Details"), 
                  command=self.view_user_details, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Edit User"), 
                  command=self.edit_user, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text=_("Disable User"), 
                  command=self.disable_user, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
    
    def load_users_data(self):
        # Clear existing data
        for item in self.users_tree.get_children():
            self.users_tree.delete(item)
        
        # In a real app, this would load users from database
        # For now, we'll show a message
        self.users_tree.insert("", tk.END, values=(
            _("N/A"), _("N/A"), _("N/A"), _("N/A"), _("N/A"), _("N/A")
        ))
    
    def add_user(self):
        # Create add user dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Add User"))
        dialog.geometry("600x700")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Create a canvas and scrollbar for the permissions section
        canvas = tk.Canvas(dialog)
        scrollbar = ttk.Scrollbar(dialog, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Form fields
        form_frame = ttk.Frame(scrollable_frame)
        form_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(form_frame, text=_("Username:")).pack(pady=5)
        username_entry = ttk.Entry(form_frame, width=40)
        username_entry.pack(pady=5)
        
        ttk.Label(form_frame, text=_("Email:")).pack(pady=5)
        email_entry = ttk.Entry(form_frame, width=40)
        email_entry.pack(pady=5)
        
        ttk.Label(form_frame, text=_("Password:")).pack(pady=5)
        password_entry = ttk.Entry(form_frame, width=40, show="*")
        password_entry.pack(pady=5)
        
        ttk.Label(form_frame, text=_("Role:")).pack(pady=5)
        role_var = tk.StringVar()
        role_combo = ttk.Combobox(form_frame, textvariable=role_var,
                                 values=[_("Admin"), _("Technician"), _("Doctor"), _("Receptionist")],
                                 state="readonly", width=37)
        role_combo.pack(pady=5)
        
        ttk.Label(form_frame, text=_("Active:")).pack(pady=5)
        active_var = tk.BooleanVar(value=True)
        active_check = ttk.Checkbutton(form_frame, variable=active_var)
        active_check.pack(pady=5)
        
        # Permissions section
        permissions_frame = ttk.LabelFrame(scrollable_frame, text=_("Permissions"), padding=10)
        permissions_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Create permission checkboxes
        permission_vars = {}
        permissions_by_category = {
            _("Patient Management"): [
                Permission.VIEW_PATIENTS, Permission.ADD_PATIENT, 
                Permission.EDIT_PATIENT, Permission.DELETE_PATIENT
            ],
            _("Test Management"): [
                Permission.VIEW_TESTS, Permission.ADD_TEST, 
                Permission.EDIT_TEST, Permission.DELETE_TEST
            ],
            _("Sample Management"): [
                Permission.VIEW_SAMPLES, Permission.ADD_SAMPLE, 
                Permission.EDIT_SAMPLE, Permission.DELETE_SAMPLE
            ],
            _("Report Management"): [
                Permission.VIEW_REPORTS, Permission.ADD_REPORT, 
                Permission.EDIT_REPORT, Permission.DELETE_REPORT, 
                Permission.SIGN_REPORT
            ],
            _("Billing Management"): [
                Permission.VIEW_BILLING, Permission.ADD_INVOICE, 
                Permission.EDIT_INVOICE, Permission.DELETE_INVOICE
            ],
            _("Inventory Management"): [
                Permission.VIEW_INVENTORY, Permission.ADD_INVENTORY, 
                Permission.EDIT_INVENTORY, Permission.DELETE_INVENTORY
            ],
            _("User Management"): [
                Permission.VIEW_USERS, Permission.ADD_USER, 
                Permission.EDIT_USER, Permission.DELETE_USER
            ],
            _("Statistics & Reports"): [
                Permission.VIEW_STATISTICS, Permission.GENERATE_REPORTS
            ]
        }
        
        for category, perms in permissions_by_category.items():
            category_frame = ttk.LabelFrame(permissions_frame, text=category)
            category_frame.pack(fill=tk.X, padx=5, pady=5)
            
            for perm in perms:
                var = tk.BooleanVar()
                permission_vars[perm] = var
                ttk.Checkbutton(
                    category_frame, 
                    text=_(perm.value), 
                    variable=var
                ).pack(anchor=tk.W, padx=5, pady=2)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Buttons frame
        button_frame = ttk.Frame(scrollable_frame)
        button_frame.pack(pady=20)
        
        def save_user():
            username = username_entry.get().strip()
            email = email_entry.get().strip()
            password = password_entry.get()
            role_text = role_var.get()
            is_active = active_var.get()
            
            # Validation
            if not username:
                messagebox.showerror(_("Error"), _("Please enter username"))
                return
            
            if not email:
                messagebox.showerror(_("Error"), _("Please enter email"))
                return
            
            if not password:
                messagebox.showerror(_("Error"), _("Please enter password"))
                return
            
            if not role_text:
                messagebox.showerror(_("Error"), _("Please select role"))
                return
            
            # Map role text to enum
            role_map = {
                _("Admin"): UserRole.ADMIN,
                _("Technician"): UserRole.TECHNICIAN,
                _("Doctor"): UserRole.DOCTOR,
                _("Receptionist"): UserRole.RECEPTIONIST
            }
            role = role_map.get(role_text, UserRole.RECEPTIONIST)
            
            # Collect selected permissions
            selected_permissions = [perm for perm, var in permission_vars.items() if var.get()]
            
            # Hash password
            password_hash = self.hash_password(password)
            
            # Create user
            user = User(
                id=str(uuid.uuid4()),
                username=username,
                email=email,
                password_hash=password_hash,
                role=role,
                is_active=is_active,
                permissions=selected_permissions
            )
            
            if self.db.create_user(user):
                messagebox.showinfo(_("Success"), _("User added successfully"))
                dialog.destroy()
                self.load_users_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to add user"))
        
        ttk.Button(button_frame, text=_("Save"), command=save_user).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), 
                  command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Focus on first entry
        username_entry.focus()

    def run_test(self):
        selected_test = self.tests_list.curselection()
        if selected_test:
            test_name = self.tests_list.get(selected_test)
            messagebox.showinfo(_("Run Test"), f"Running test: {test_name}")

    def view_user_details(self):
        selected = self.users_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a user"))
            return
        
        # Get the selected user
        item = self.users_tree.item(selected[0])
        values = item['values']
        user_id = values[0]
        
        # Get user from database
        user = self.db.get_user(user_id)
        if not user:
            messagebox.showerror(_("Error"), _("User not found"))
            return
        
        # Create detail dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("User Details"))
        dialog.geometry("600x700")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Create a canvas and scrollbar for the permissions section
        canvas = tk.Canvas(dialog)
        scrollbar = ttk.Scrollbar(dialog, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # User info
        info_frame = ttk.LabelFrame(scrollable_frame, text=_("User Information"), padding=10)
        info_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(info_frame, text=f"{_('Username')}: {user.username}").pack(anchor=tk.W)
        ttk.Label(info_frame, text=f"{_('Email')}: {user.email}").pack(anchor=tk.W)
        ttk.Label(info_frame, text=f"{_('Role')}: {_(user.role.value)}").pack(anchor=tk.W)
        ttk.Label(info_frame, text=f"{_('Active')}: {_('Yes') if user.is_active else _('No')}").pack(anchor=tk.W)
        ttk.Label(info_frame, text=f"{_('Created At')}: {user.created_at.strftime('%Y-%m-%d %H:%M')}").pack(anchor=tk.W)
        if user.last_login:
            ttk.Label(info_frame, text=f"{_('Last Login')}: {user.last_login.strftime('%Y-%m-%d %H:%M')}").pack(anchor=tk.W)
        
        # Permissions section
        permissions_frame = ttk.LabelFrame(scrollable_frame, text=_("Permissions"), padding=10)
        permissions_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        if user.permissions:
            # Group permissions by category for display
            permissions_by_category = {
                _("Patient Management"): [],
                _("Test Management"): [],
                _("Sample Management"): [],
                _("Report Management"): [],
                _("Billing Management"): [],
                _("Inventory Management"): [],
                _("User Management"): [],
                _("Statistics & Reports"): []
            }
            
            # Categorize permissions
            for perm in user.permissions:
                if perm in [Permission.VIEW_PATIENTS, Permission.ADD_PATIENT, 
                           Permission.EDIT_PATIENT, Permission.DELETE_PATIENT]:
                    permissions_by_category[_("Patient Management")].append(perm)
                elif perm in [Permission.VIEW_TESTS, Permission.ADD_TEST, 
                             Permission.EDIT_TEST, Permission.DELETE_TEST]:
                    permissions_by_category[_("Test Management")].append(perm)
                elif perm in [Permission.VIEW_SAMPLES, Permission.ADD_SAMPLE, 
                             Permission.EDIT_SAMPLE, Permission.DELETE_SAMPLE]:
                    permissions_by_category[_("Sample Management")].append(perm)
                elif perm in [Permission.VIEW_REPORTS, Permission.ADD_REPORT, 
                             Permission.EDIT_REPORT, Permission.DELETE_REPORT, 
                             Permission.SIGN_REPORT]:
                    permissions_by_category[_("Report Management")].append(perm)
                elif perm in [Permission.VIEW_BILLING, Permission.ADD_INVOICE, 
                             Permission.EDIT_INVOICE, Permission.DELETE_INVOICE]:
                    permissions_by_category[_("Billing Management")].append(perm)
                elif perm in [Permission.VIEW_INVENTORY, Permission.ADD_INVENTORY, 
                             Permission.EDIT_INVENTORY, Permission.DELETE_INVENTORY]:
                    permissions_by_category[_("Inventory Management")].append(perm)
                elif perm in [Permission.VIEW_USERS, Permission.ADD_USER, 
                             Permission.EDIT_USER, Permission.DELETE_USER]:
                    permissions_by_category[_("User Management")].append(perm)
                elif perm in [Permission.VIEW_STATISTICS, Permission.GENERATE_REPORTS]:
                    permissions_by_category[_("Statistics & Reports")].append(perm)
            
            # Display permissions by category
            for category, perms in permissions_by_category.items():
                if perms:
                    category_frame = ttk.LabelFrame(permissions_frame, text=category)
                    category_frame.pack(fill=tk.X, padx=5, pady=5)
                    
                    for perm in perms:
                        ttk.Label(category_frame, text=_(perm.value)).pack(anchor=tk.W, padx=5, pady=2)
        else:
            ttk.Label(permissions_frame, text=_("No permissions assigned")).pack()
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Close button
        button_frame = ttk.Frame(scrollable_frame)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(button_frame, text=_("Close"), command=dialog.destroy).pack(side=tk.RIGHT)

    def edit_user(self):
        selected = self.users_tree.selection()
        if not selected:
            messagebox.showwarning(_("Warning"), _("Please select a user to edit"))
            return
        
        # Get the selected user
        item = self.users_tree.item(selected[0])
        values = item['values']
        user_id = values[0]
        
        # Get user from database
        user = self.db.get_user(user_id)
        if not user:
            messagebox.showerror(_("Error"), _("User not found"))
            return
        
        # Create edit user dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(_("Edit User"))
        dialog.geometry("600x700")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Create a canvas and scrollbar for the permissions section
        canvas = tk.Canvas(dialog)
        scrollbar = ttk.Scrollbar(dialog, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Form fields
        form_frame = ttk.Frame(scrollable_frame)
        form_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(form_frame, text=_("Username:")).pack(pady=5)
        username_entry = ttk.Entry(form_frame, width=40)
        username_entry.pack(pady=5)
        username_entry.insert(0, user.username)
        
        ttk.Label(form_frame, text=_("Email:")).pack(pady=5)
        email_entry = ttk.Entry(form_frame, width=40)
        email_entry.pack(pady=5)
        email_entry.insert(0, user.email)
        
        ttk.Label(form_frame, text=_("Role:")).pack(pady=5)
        role_var = tk.StringVar()
        role_combo = ttk.Combobox(form_frame, textvariable=role_var,
                                 values=[_("Admin"), _("Technician"), _("Doctor"), _("Receptionist")],
                                 state="readonly", width=37)
        role_combo.pack(pady=5)
        
        # Set role combo value
        role_text_map = {
            UserRole.ADMIN: _("Admin"),
            UserRole.TECHNICIAN: _("Technician"),
            UserRole.DOCTOR: _("Doctor"),
            UserRole.RECEPTIONIST: _("Receptionist")
        }
        role_combo.set(role_text_map.get(user.role, _("Receptionist")))
        
        ttk.Label(form_frame, text=_("Active:")).pack(pady=5)
        active_var = tk.BooleanVar(value=user.is_active)
        active_check = ttk.Checkbutton(form_frame, variable=active_var)
        active_check.pack(pady=5)
        
        # Permissions section
        permissions_frame = ttk.LabelFrame(scrollable_frame, text=_("Permissions"), padding=10)
        permissions_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Create permission checkboxes
        permission_vars = {}
        permissions_by_category = {
            _("Patient Management"): [
                Permission.VIEW_PATIENTS, Permission.ADD_PATIENT, 
                Permission.EDIT_PATIENT, Permission.DELETE_PATIENT
            ],
            _("Test Management"): [
                Permission.VIEW_TESTS, Permission.ADD_TEST, 
                Permission.EDIT_TEST, Permission.DELETE_TEST
            ],
            _("Sample Management"): [
                Permission.VIEW_SAMPLES, Permission.ADD_SAMPLE, 
                Permission.EDIT_SAMPLE, Permission.DELETE_SAMPLE
            ],
            _("Report Management"): [
                Permission.VIEW_REPORTS, Permission.ADD_REPORT, 
                Permission.EDIT_REPORT, Permission.DELETE_REPORT, 
                Permission.SIGN_REPORT
            ],
            _("Billing Management"): [
                Permission.VIEW_BILLING, Permission.ADD_INVOICE, 
                Permission.EDIT_INVOICE, Permission.DELETE_INVOICE
            ],
            _("Inventory Management"): [
                Permission.VIEW_INVENTORY, Permission.ADD_INVENTORY, 
                Permission.EDIT_INVENTORY, Permission.DELETE_INVENTORY
            ],
            _("User Management"): [
                Permission.VIEW_USERS, Permission.ADD_USER, 
                Permission.EDIT_USER, Permission.DELETE_USER
            ],
            _("Statistics & Reports"): [
                Permission.VIEW_STATISTICS, Permission.GENERATE_REPORTS
            ]
        }
        
        for category, perms in permissions_by_category.items():
            category_frame = ttk.LabelFrame(permissions_frame, text=category)
            category_frame.pack(fill=tk.X, padx=5, pady=5)
            
            for perm in perms:
                var = tk.BooleanVar(value=perm in user.permissions)
                permission_vars[perm] = var
                ttk.Checkbutton(
                    category_frame, 
                    text=_(perm.value), 
                    variable=var
                ).pack(anchor=tk.W, padx=5, pady=2)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Buttons frame
        button_frame = ttk.Frame(scrollable_frame)
        button_frame.pack(pady=20)
        
        def save_user():
            username = username_entry.get().strip()
            email = email_entry.get().strip()
            role_text = role_var.get()
            is_active = active_var.get()
            
            # Validation
            if not username:
                messagebox.showerror(_("Error"), _("Please enter username"))
                return
            
            if not email:
                messagebox.showerror(_("Error"), _("Please enter email"))
                return
            
            if not role_text:
                messagebox.showerror(_("Error"), _("Please select role"))
                return
            
            # Map role text to enum
            role_map = {
                _("Admin"): UserRole.ADMIN,
                _("Technician"): UserRole.TECHNICIAN,
                _("Doctor"): UserRole.DOCTOR,
                _("Receptionist"): UserRole.RECEPTIONIST
            }
            role = role_map.get(role_text, UserRole.RECEPTIONIST)
            
            # Collect selected permissions
            selected_permissions = [perm for perm, var in permission_vars.items() if var.get()]
            
            # Update user object
            user.username = username
            user.email = email
            user.role = role
            user.is_active = is_active
            user.permissions = selected_permissions
            
            if self.db.update_user(user):
                messagebox.showinfo(_("Success"), _("User updated successfully"))
                dialog.destroy()
                self.load_users_data()
            else:
                messagebox.showerror(_("Error"), _("Failed to update user"))
        
        ttk.Button(button_frame, text=_("Save"), command=save_user).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text=_("Cancel"), 
                  command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Focus on first entry
        username_entry.focus()

    def disable_user(self):
        messagebox.showinfo(_("Disable User"), _("User disabling functionality would be implemented here"))
    
    def show_statistics(self):
        self.current_screen = self.show_statistics
        self.clear_content()
        # Statistics header with 3D styling
        header_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        header_frame.pack(fill=tk.X, padx=10, pady=10)
        ttk.Label(header_frame, text=_("Reports and Statistics"), style="Title.TLabel").pack(side=tk.LEFT)
        # Date range selector
        date_frame = ttk.Frame(header_frame, style="Card.TFrame")
        date_frame.pack(side=tk.RIGHT)
        ttk.Label(date_frame, text=_("From:"), font=("Arial", 10), foreground="#000080").pack(side=tk.LEFT)
        from_date_entry = ttk.Entry(date_frame, width=10)
        from_date_entry.pack(side=tk.LEFT, padx=5)
        ttk.Label(date_frame, text=_("To:"), font=("Arial", 10), foreground="#000080").pack(side=tk.LEFT)
        to_date_entry = ttk.Entry(date_frame, width=10)
        to_date_entry.pack(side=tk.LEFT, padx=5)
        def on_generate():
            self.generate_statistics(
                from_date=from_date_entry.get(),
                to_date=to_date_entry.get(),
                patient_frame=patient_frame,
                financial_frame=financial_frame
            )
        ttk.Button(date_frame, text=_("Generate"), command=on_generate, style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        # Statistics panels with 3D styling
        stats_frame = ttk.Frame(self.content_frame, style="Card.TFrame")
        stats_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        # Create notebook for different report types
        notebook = ttk.Notebook(stats_frame)
        notebook.pack(fill=tk.BOTH, expand=True)
        # Patient statistics tab
        patient_frame = ttk.Frame(notebook, style="Card.TFrame")
        notebook.add(patient_frame, text=_("Patient Statistics"))
        # Test statistics tab (unchanged for now)
        test_frame = ttk.Frame(notebook, style="Card.TFrame")
        notebook.add(test_frame, text=_("Test Statistics"))
        # Financial statistics tab
        financial_frame = ttk.Frame(notebook, style="Card.TFrame")
        notebook.add(financial_frame, text=_("Financial Statistics"))
        # Inventory statistics tab (unchanged for now)
        inventory_frame = ttk.Frame(notebook, style="Card.TFrame")
        notebook.add(inventory_frame, text=_("Inventory Statistics"))
        # Initial statistics
        self.generate_statistics(None, None, patient_frame, financial_frame)
        # Add report generation buttons
        ttk.Button(patient_frame, text=_("Generate Patient Report"), command=self.generate_patient_report, style="Accent.TButton").pack(pady=10)
        ttk.Button(financial_frame, text=_("Generate Financial Report"), command=self.generate_financial_report, style="Accent.TButton").pack(pady=10)
    
    def generate_statistics(self, from_date, to_date, patient_frame, financial_frame):
        # Clear previous stats
        for widget in patient_frame.winfo_children():
            widget.destroy()
        for widget in financial_frame.winfo_children():
            widget.destroy()
        # Parse dates if provided
        from_dt = None
        to_dt = None
        from datetime import datetime
        try:
            if from_date:
                from_dt = datetime.strptime(from_date, "%Y-%m-%d")
            if to_date:
                to_dt = datetime.strptime(to_date, "%Y-%m-%d")
        except Exception:
            pass
        # Patient statistics
        patients = self.db.get_all_patients()
        total_patients = len(patients)
        new_patients = 0
        returning_patients = 0
        now = datetime.now()
        for p in patients:
            if from_dt and to_dt:
                if from_dt <= p.created_at <= to_dt:
                    new_patients += 1
            elif p.created_at.month == now.month and p.created_at.year == now.year:
                new_patients += 1
        returning_patients = total_patients - new_patients
        ttk.Label(patient_frame, text=_("Total Patients: {}".format(total_patients)), font=("Arial", 12, "bold"), foreground="#000080").pack(pady=10)
        ttk.Label(patient_frame, text=_("New Patients (This Month): {}".format(new_patients)), font=("Arial", 12, "bold"), foreground="#000080").pack(pady=10)
        ttk.Label(patient_frame, text=_("Returning Patients: {}".format(returning_patients)), font=("Arial", 12, "bold"), foreground="#000080").pack(pady=10)
        # Financial statistics
        from collections import defaultdict
        test_types = {t.id: t for t in self.db.get_all_test_types()}
        test_requests = self.db.get_all_test_requests()
        total_revenue = 0
        outstanding_payments = 0
        test_prices = []
        for tr in test_requests:
            test_type = test_types.get(tr.test_type_id)
            if test_type:
                test_prices.append(test_type.price)
                total_revenue += test_type.price
        # Outstanding payments (if you have payment tracking, update here)
        # For now, assume all are paid
        avg_test_price = (sum(test_prices) / len(test_prices)) if test_prices else 0
        ttk.Label(financial_frame, text=_("Total Revenue: ${:,.2f}".format(total_revenue)), font=("Arial", 12, "bold"), foreground="#000080").pack(pady=10)
        ttk.Label(financial_frame, text=_("Outstanding Payments: ${:,.2f}".format(outstanding_payments)), font=("Arial", 12, "bold"), foreground="#000080").pack(pady=10)
        ttk.Label(financial_frame, text=_("Average Test Price: ${:,.2f}".format(avg_test_price)), font=("Arial", 12, "bold"), foreground="#000080").pack(pady=10)

    def generate_patient_report(self):
        import tkinter.filedialog as fd
        from tkinter import messagebox
        patients = self.db.get_all_patients()
        filetypes = [("CSV Files", "*.csv"), ("PDF Files", "*.pdf")]
        file_path = fd.asksaveasfilename(defaultextension=".csv", filetypes=filetypes)
        if not file_path:
            return
        if file_path.endswith('.pdf'):
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib import colors
            from reportlab.platypus import Table, TableStyle, SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            data = [["ID", "Name", "Age", "Gender", "Contact", "Created At"]]
            for p in patients:
                data.append([p.id, p.name, p.age, p.gender.value, p.contact_info, p.created_at.strftime("%Y-%m-%d %H:%M:%S")])
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            styles = getSampleStyleSheet()
            elements = [Paragraph("Patient Statistics Report", styles['Title']), Spacer(1, 12)]
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.lightblue),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0,0), (-1,0), 12),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
            ]))
            elements.append(table)
            doc.build(elements)
            messagebox.showinfo(_("Success"), _("Patient statistics PDF exported successfully."))
        else:
            import csv
            with open(file_path, "w", newline="") as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow(["ID", "Name", "Age", "Gender", "Contact", "Created At"])
                for p in patients:
                    writer.writerow([p.id, p.name, p.age, p.gender.value, p.contact_info, p.created_at.strftime("%Y-%m-%d %H:%M:%S")])
            messagebox.showinfo(_("Success"), _("Patient statistics report exported successfully."))

    def generate_financial_report(self):
        import tkinter.filedialog as fd
        from tkinter import messagebox
        test_types = {t.id: t for t in self.db.get_all_test_types()}
        test_requests = self.db.get_all_test_requests()
        filetypes = [("CSV Files", "*.csv"), ("PDF Files", "*.pdf")]
        file_path = fd.asksaveasfilename(defaultextension=".csv", filetypes=filetypes)
        if not file_path:
            return
        if file_path.endswith('.pdf'):
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib import colors
            from reportlab.platypus import Table, TableStyle, SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            data = [["Test Request ID", "Patient ID", "Test Type", "Price", "Requested At"]]
            for tr in test_requests:
                test_type = test_types.get(tr.test_type_id)
                data.append([
                    tr.id, tr.patient_id, test_type.name if test_type else "", test_type.price if test_type else 0, tr.requested_at.strftime("%Y-%m-%d %H:%M:%S")
                ])
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            styles = getSampleStyleSheet()
            elements = [Paragraph("Financial Report", styles['Title']), Spacer(1, 12)]
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.lightblue),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0,0), (-1,0), 12),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
            ]))
            elements.append(table)
            doc.build(elements)
            messagebox.showinfo(_("Success"), _("Financial PDF report exported successfully."))
        else:
            import csv
            with open(file_path, "w", newline="") as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow(["Test Request ID", "Patient ID", "Test Type", "Price", "Requested At"])
                for tr in test_requests:
                    test_type = test_types.get(tr.test_type_id)
                    writer.writerow([
                        tr.id, tr.patient_id, test_type.name if test_type else "", test_type.price if test_type else 0, tr.requested_at.strftime("%Y-%m-%d %H:%M:%S")
                    ])
            messagebox.showinfo(_("Success"), _("Financial report exported successfully."))
    
    def clear_content(self):
        for widget in self.content_frame.winfo_children():
            widget.destroy()

def main():
    root = tk.Tk()
    app = MedicalLabApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
